import pdfplumber
import re
import os
import uuid
import tempfile
from flask import Flask, request, render_template, redirect, url_for, session, send_file
from werkzeug.utils import secure_filename
import json
import openpyxl
from datetime import datetime, timedelta
import zipfile
import shutil

# --- Import para manipulação de DOCX ---
from docx import Document
from docx.shared import Pt
from docx.text.run import Run

# --- Import para manipulação de imagens e PDF ---
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, portrait
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader

# --- FUNÇÃO AUXILIAR GLOBAL PARA FORMATAR VALORES PARA EXIBIÇÃO ---
def format_value_for_display(value, is_date=False, is_numeric=False, date_separator='/', numeric_decimal_separator=','):
    if value is None or value == 'Não informado' or value == 'Não calculado' or str(value).strip() == '':
        return 'Não informado'

    s_value = str(value).strip()

    if is_date:
        try: # Tenta YYYY-MM-DD
            dt_obj = datetime.strptime(s_value, '%Y-%m-%d')
            return dt_obj.strftime(f'%d{date_separator}%m{date_separator}%Y')
        except ValueError:
            try: # Tenta DD/MM/YYYY
                dt_obj = datetime.strptime(s_value, '%d/%m/%Y')
                return dt_obj.strftime(f'%d{date_separator}%m{date_separator}%Y')
            except ValueError:
                try: # Tenta DD-MM-YYYY
                    dt_obj = datetime.strptime(s_value, '%d-%m-%Y')
                    return dt_obj.strftime(f'%d{date_separator}%m{date_separator}%Y')
                except ValueError:
                    return s_value
    
    if is_numeric:
        try:
            f_value = float(s_value.replace(',', '.'))
            if f_value == int(f_value):
                return str(int(f_value))
            else:
                formatted_str = f" {f_value:.3f}".rstrip('0')
                if formatted_str.endswith('.'):
                    formatted_str = formatted_str.rstrip('.')
                return formatted_str.replace('.', numeric_decimal_separator)
        except ValueError:
            return s_value

    return s_value

# --- Funções Auxiliares de Extração para Layouts Específicos ---

def _extrair_dados_layout_adriano_style(texto):
    dados_extraidos = {
        'Nome_Razao_Social': 'Não encontrado', 'Endereco_Rua_Numero': 'Não encontrado',
        'Bairro': 'Não encontrado', 'CNPJ_CPF': 'Não encontrado',
        'Cidade': 'Não encontrado', 'CEP': 'Não encontrado',
        'UC': 'Não encontrado', 'Grupo_Tarifario': 'Não encontrado',
        'Tensao_Nominal_V': 'Não encontrado', 'Classe_Tarifaria': 'Não encontrado',
        'Estado': 'Não encontrado',
    }

    try:
        match_tensao = re.search(r'TENSÃO NOMINAL EM VOLTS\s*Disp\.:\s*(\d+)', texto)
        if match_tensao:
            dados_extraidos['Tensao_Nominal_V'] = int(match_tensao.group(1))

        match_nome = re.search(r'Inscrição no CNPJ: \d{2}\.\d{3}\.\d{3}\/\d{4}-\d{2}\n+([A-Z\s,.]+)\n', texto)
        customer_name_found = None
        if match_nome:
            customer_name_found = match_nome.group(1).strip()
            dados_extraidos['Nome_Razao_Social'] = customer_name_found

        if customer_name_found and customer_name_found != 'Não encontrado':
            street_and_number_pattern = r'((?:R|AV|EST|ROD|AL|TV|PR|TR|VD|RUA|VL|PRC|PCA)\s+[A-Z\s,.-]+?\s*\d+\s*(?:[A-Z0-9\s,.-]+)?)'
            address_block_full_regex = (
                re.escape(customer_name_found) + r'.*?' 
                + street_and_number_pattern + r'\n' 
                + r'([A-Z\s,.-]+)\n' 
                + r'(\d{5}-\d{3})\s+([A-Z\s,.-]+)\s+(RS)'
            )
            match_endereco_bloco = re.search(address_block_full_regex, texto, re.DOTALL)
            if match_endereco_bloco:
                dados_extraidos['Endereco_Rua_Numero'] = match_endereco_bloco.group(1).strip()
                dados_extraidos['Bairro'] = match_endereco_bloco.group(2).strip()
                dados_extraidos['CEP'] = match_endereco_bloco.group(3).strip()
                dados_extraidos['Cidade'] = match_endereco_bloco.group(4).strip()
                dados_extraidos['Estado'] = match_endereco_bloco.group(5).strip()

        match_cpf = re.search(r'CPF:\s*(\d{3}\.\d{3}\.\d{3}-\d{2})', texto)
        if match_cpf:
            dados_extraidos['CNPJ_CPF'] = match_cpf.group(1)
        else:
            match_cnpj = re.search(r'CNPJ:\s*(\d{2}\.\d{3}\.\d{3}\/\d{4}-\d{2})', texto)
            if match_cnpj:
                dados_extraidos['CNPJ_CPF'] = match_cnpj.group(1)

        match_uc = re.search(r'UC:\s*(\d{10})', texto)
        if match_uc:
            dados_extraidos['UC'] = match_uc.group(1)
        else:
            match_uc_alt = re.search(r'Lim.\s*máx.:\s*\d+\s*(\d{10})', texto)
            if match_uc_alt:
                dados_extraidos['UC'] = match_uc_alt.group(1)

        match_classificacao = re.search(r'Classificaç(?:ão|ao):\s*([^\n]+)', texto, re.IGNORECASE)
        if match_classificacao:
            classif = match_classificacao.group(1).strip().replace('Tipo de Fornecimento:', '').strip()
            match_grupo = re.search(r'(B[1-4]|A)', classif)
            if match_grupo:
                dados_extraidos['Grupo_Tarifario'] = match_grupo.group(1)
            match_classe = re.search(r'(Residencial|Comercial|Industrial|Rural|Poder Público|Iluminação Pública)', classif, re.IGNORECASE)
            if match_classe:
                dados_extraidos['Classe_Tarifaria'] = match_classe.group(1)

    except Exception as e:
        print(f"Erro ao extrair dados do layout 'Adriano Style': {e}")
        pass

    return dados_extraidos


def _extrair_dados_layout_adroaldo_style(texto):
    dados_extraidos = {
        'Nome_Razao_Social': 'Não encontrado', 'Endereco_Rua_Numero': 'Não encontrado',
        'Bairro': 'Não encontrado', 'CNPJ_CPF': 'Não encontrado',
        'Cidade': 'Não encontrado', 'CEP': 'Não encontrado',
        'UC': 'Não encontrado', 'Grupo_Tarifario': 'Não encontrado',
        'Tensao_Nominal_V': 'Não encontrado', 'Classe_Tarifaria': 'Não encontrado',
        'Estado': 'Não encontrado',
    }

    try:
        match_tensao = re.search(r'TENSÃO NOMINAL EM VOLTS\s*Disp\.:\s*(\d+)', texto)
        if match_tensao:
            dados_extraidos['Tensao_Nominal_V'] = int(match_tensao.group(1))

        match_nome = re.search(r'Inscrição no CNPJ: \d{2}\.\d{3}\.\d{3}\/\d{4}-\d{2}\n+([A-Z\s,.]+)\n', texto)
        customer_name_found = None
        if match_nome:
            customer_name_found = match_nome.group(1).strip()
            dados_extraidos['Nome_Razao_Social'] = customer_name_found

        if customer_name_found and customer_name_found != 'Não encontrado':
            street_and_number_pattern = r'((?:R|AV|EST|ROD|AL|TV|PR|TR|VD|RUA|VL|PRC|PCA)\s+[A-Z\s,.-]+?\s*\d+\s*(?:[A-Z0-9\s,.-]+)?)'
            address_block_full_regex = (
                re.escape(customer_name_found) + r'.*?' 
                + street_and_number_pattern + r'\n' 
                + r'([A-Z\s,.-]+)\n' 
                + r'(\d{5}-\d{3})\s+([A-Z\s,.-]+)\s+(RS)'
            )
            match_endereco_bloco = re.search(address_block_full_regex, texto, re.DOTALL)
            if match_endereco_bloco:
                dados_extraidos['Endereco_Rua_Numero'] = match_endereco_bloco.group(1).strip()
                dados_extraidos['Bairro'] = match_endereco_bloco.group(2).strip()
                dados_extraidos['CEP'] = match_endereco_bloco.group(3).strip()
                dados_extraidos['Cidade'] = match_endereco_bloco.group(4).strip()
                dados_extraidos['Estado'] = match_endereco_bloco.group(5).strip()

        match_cpf_masked = re.search(r'CPF:\s*(((\*)?){6}\.\d{3}-(((\*)?){2}))', texto)
        if match_cpf_masked:
            dados_extraidos['CNPJ_CPF'] = match_cpf_masked.group(1)
        else:
            match_cnpj = re.search(r'CNPJ:\s*(\d{2}\.\d{3}\.\d{3}\/\d{4}-\d{2})', texto)
            if match_cnpj:
                dados_extraidos['CNPJ_CPF'] = match_cnpj.group(1)

        match_uc = re.search(r'Lim.\s*máx.:\s*\d+\s*(\d{10})', texto)
        if match_uc:
            dados_extraidos['UC'] = match_uc.group(1)

        match_classificacao = re.search(r'Classificaç(?:ão|ao):\s*([^\n]+)', texto, re.IGNORECASE)
        if match_classificacao:
            classif = match_classificacao.group(1).strip().replace('Tipo de Fornecimento:', '').strip()
            match_grupo = re.search(r'(B[1-4]|A)', classif)
            if match_grupo:
                dados_extraidos['Grupo_Tarifario'] = match_grupo.group(1)
            match_classe = re.search(r'(Residencial|Comercial|Industrial|Rural|Poder Público|Iluminação Pública)', classif, re.IGNORECASE)
            if match_classe:
                dados_extraidos['Classe_Tarifaria'] = match_classe.group(1)

    except Exception as e:
        print(f"Erro ao extrair dados do layout 'Adroaldo Style': {e}")
        pass

    return dados_extraidos


def _extrair_dados_layout_arcindo_style(texto):
    dados_extraidos = {
        'Nome_Razao_Social': 'Não encontrado', 'Endereco_Rua_Numero': 'Não encontrado',
        'Bairro': 'Não encontrado', 'CNPJ_CPF': 'Não encontrado',
        'Cidade': 'Não encontrado', 'CEP': 'Não encontrado',
        'UC': 'Não encontrado', 'Grupo_Tarifario': 'Não encontrado',
        'Tensao_Nominal_V': 'Não encontrado', 'Classe_Tarifaria': 'Não encontrado',
        'Estado': 'Não encontrado',
    }

    try:
        match_tensao = re.search(r'TENSÃO NOMINAL EM VOLTS\s*Disp\.:\s*(\d+)', texto)
        if match_tensao:
            dados_extraidos['Tensao_Nominal_V'] = int(match_tensao.group(1))

        match_nome = re.search(r'CÓDIGO DA UNIDADE CONSUMIDORA:\s*\d+\n([A-Z\s]+)\n', texto)
        customer_name_found = None
        if match_nome:
            customer_name_found = match_nome.group(1).strip()
            dados_extraidos['Nome_Razao_Social'] = customer_name_found
        
        if customer_name_found and customer_name_found != 'Não encontrado':
            street_and_number_pattern = r'((?:R|AV|EST|ROD|AL|TV|PR|TR|VD|RUA|VL|PRC|PCA)\s+[A-Z\s,.-]+?\s*\d+\s*(?:[A-Z0-9\s,.-]+)?)'
            address_block_full_regex = (
                re.escape(customer_name_found) + r'.*?' 
                + street_and_number_pattern + r'\n' 
                + r'([A-Z\s,.-]+)\n' 
                + r'(\d{5}-\d{3})\s+([A-Z\s,.-]+)\s*-\s*(RS)'
            )
            match_endereco_bloco = re.search(address_block_full_regex, texto, re.DOTALL)
            if match_endereco_bloco:
                dados_extraidos['Endereco_Rua_Numero'] = match_endereco_bloco.group(1).strip()
                dados_extraidos['Bairro'] = match_endereco_bloco.group(2).strip()
                dados_extraidos['CEP'] = match_endereco_bloco.group(3).strip()
                dados_extraidos['Cidade'] = match_endereco_bloco.group(4).strip()
                dados_extraidos['Estado'] = match_endereco_bloco.group(5).strip()

        match_cpf_masked = re.search(r'CPF:\s*(((\*)?){6}\.\d{3}-(((\*)?){2}))', texto)
        if match_cpf_masked:
            dados_extraidos['CNPJ_CPF'] = match_cpf_masked.group(1)
        else:
            match_cnpj = re.search(r'CNPJ:\s*(\d{2}\.\d{3}\.\d{3}\/\d{4}-\d{2})', texto)
            if match_cnpj:
                dados_extraidos['CNPJ_CPF'] = match_cnpj.group(1)

        match_uc = re.search(r'CÓDIGO DA UNIDADE CONSUMIDORA:\s*(\d{10})', texto)
        if match_uc:
            dados_extraidos['UC'] = match_uc.group(1)
        else:
            match_uc_alt = re.search(r'(\d{10})\n1/2', texto)
            if match_uc_alt:
                dados_extraidos['UC'] = match_uc_alt.group(1)

        match_classificacao = re.search(r'Classificaç(?:ão|ao):\s*([^\n]+)', texto, re.IGNORECASE)
        if match_classificacao:
            classif = match_classificacao.group(1).strip()
            match_grupo = re.search(r'(B[1-4]|A)', classif)
            if match_grupo:
                dados_extraidos['Grupo_Tarifario'] = match_grupo.group(1)
            match_classe = re.search(r'(Residencial|Comercial|Industrial|Rural|Poder Público|Iluminação Pública)', classif, re.IGNORECASE)
            if match_classe:
                dados_extraidos['Classe_Tarifaria'] = match_classe.group(1)

    except Exception as e:
        print(f"Erro ao extrair dados do layout 'Arcindo Style': {e}")
        pass

    return dados_extraidos

# --- Funções para Sub-layouts da Cooperluz ---

def _extrair_dados_layout_cooperluz_sublayout_com_cod_ua(texto):
    dados_extraidos = {
        'Nome_Razao_Social': 'Não encontrado', 'Endereco_Rua_Numero': 'Não encontrado',
        'Bairro': 'Não encontrado', 'CNPJ_CPF': 'Não encontrado',
        'Cidade': 'Não encontrado', 'CEP': 'Não encontrado',
        'UC': 'Não encontrado', 'Grupo_Tarifario': 'Não encontrado',
        'Tensao_Nominal_V': 'Não encontrado', 'Classe_Tarifaria': 'Não encontrado',
        'Estado': 'Não encontrado',
    }
    
    try:
        tipo_fornecimento_match = re.search(r'Tipo de Fornecimento:\s*(?:[\s\S]*?)(Monofásico|Bifásico|Trifásico)', texto, re.IGNORECASE | re.DOTALL)
        if tipo_fornecimento_match:
            tipo_fornecimento_extraido = tipo_fornecimento_match.group(1).strip()
            if 'Bifásico' in tipo_fornecimento_extraido or 'Monofásico' in tipo_fornecimento_extraido:
                dados_extraidos['Tensao_Nominal_V'] = 220
            elif 'Trifásico' in tipo_fornecimento_extraido:
                dados_extraidos['Tensao_Nominal_V'] = 380

        classificacao_line_match = re.search(r'Classificaç(?:ão|ao):\s*(.*?)(?:(?=\nTipo de Fornecimento)|\n|$)', texto, re.DOTALL | re.IGNORECASE)
        if classificacao_line_match:
            classif_line_content = classificacao_line_match.group(1).strip()
            match_grupo = re.search(r'(B[1-4]|A)', classif_line_content)
            if match_grupo:
                dados_extraidos['Grupo_Tarifario'] = match_grupo.group(1)
            match_classe = re.search(r'(Residencial|Comercial|Industrial|Rural|Poder Público|Iluminação Pública)', classif_line_content, re.IGNORECASE)
            if match_classe:
                dados_extraidos['Classe_Tarifaria'] = match_classe.group(1)

        nome_match = re.search(r'(?:Monofásico|Bifásico|Trifásico)\n+([A-Z\s,.-]+)\n+(?:Leitura anterior|DATAS DE|COD UA)', texto, re.DOTALL | re.IGNORECASE)
        if nome_match:
            dados_extraidos['Nome_Razao_Social'] = nome_match.group(1).strip()

        endereco_rua_match = re.search(r'Proxima Leitura\n+([^\n]+) DATAS DE', texto, re.DOTALL) 
        if endereco_rua_match:
            dados_extraidos['Endereco_Rua_Numero'] = endereco_rua_match.group(1).strip()

        interior_line_match = re.search(r'COD UA \d+ LEITURAS.*?\n\s*(INTERIOR / ([A-Za-zÀ-ÖØ-öø-ÿ\s,.-]+))-([A-Z]{2})', texto, re.DOTALL)
        if interior_line_match:
            dados_extraidos['Bairro'] = 'INTERIOR'
            dados_extraidos['Cidade'] = interior_line_match.group(2).strip()
            dados_extraidos['Estado'] = interior_line_match.group(3).strip()

        match_cpf_cnpj = re.search(r'CPF/CNPJ:\s*([\d*]{3}\.[\d*]{3}\.[\d*]{3}-\d{2}|\d{2}\.[\d*]{3}\.[\d*]{3}\/\d{4}-\d{2})', texto)
        if match_cpf_cnpj:
            dados_extraidos['CNPJ_CPF'] = match_cpf_cnpj.group(1)

        match_cep = re.search(r'CEP:\s*(\d{2}\s*\d{3}-\d{3})', texto)
        if match_cep:
            dados_extraidos['CEP'] = match_cep.group(1)

        uc_match = re.search(r'CEP:\s*\d{2}\s*\d{3}-\d{3}\s*([\d-]+)', texto)
        if uc_match:
            dados_extraidos['UC'] = uc_match.group(1)

    except Exception as e:
        print(f"Erro ao extrair dados do sub-layout 'Cooperluz (com COD UA)': {e}")
        pass
    return dados_extraidos

def _extrair_dados_layout_cooperluz_sublayout_sem_cod_ua(texto):
    dados_extraidos = {
        'Nome_Razao_Social': 'Não encontrado', 'Endereco_Rua_Numero': 'Não encontrado',
        'Bairro': 'Não encontrado', 'CNPJ_CPF': 'Não encontrado',
        'Cidade': 'Não encontrado', 'CEP': 'Não encontrado',
        'UC': 'Não encontrado', 'Grupo_Tarifario': 'Não encontrado',
        'Tensao_Nominal_V': 'Não encontrado', 'Classe_Tarifaria': 'Não encontrado',
        'Estado': 'Não encontrado',
    }

    try:
        tipo_fornecimento_match = re.search(r'Tipo de Fornecimento:\s*(?:[\s\S]*?)(Monofásico|Bifásico|Trifásico)', texto, re.IGNORECASE | re.DOTALL)
        if tipo_fornecimento_match:
            tipo_fornecimento_extraido = tipo_fornecimento_match.group(1).strip()
            if 'Bifásico' in tipo_fornecimento_extraido or 'Monofásico' in tipo_fornecimento_extraido:
                dados_extraidos['Tensao_Nominal_V'] = 220
            elif 'Trifásico' in tipo_fornecimento_extraido:
                dados_extraidos['Tensao_Nominal_V'] = 380

        classificacao_line_match = re.search(r'Classificaç(?:ão|ao):\s*(.*?)(?:(?=\nTipo de Fornecimento)|\n|$)', texto, re.DOTALL | re.IGNORECASE)
        if classificacao_line_match:
            classif_line_content = classificacao_line_match.group(1).strip()
            match_grupo = re.search(r'(B[1-4]|A)', classif_line_content)
            if match_grupo:
                dados_extraidos['Grupo_Tarifario'] = match_grupo.group(1)
            match_classe = re.search(r'(Residencial|Comercial|Industrial|Rural|Poder Público|Iluminação Pública)', classif_line_content, re.IGNORECASE)
            if match_classe:
                dados_extraidos['Classe_Tarifaria'] = match_classe.group(1)

        nome_match = re.search(r'(?:Monofásico|Bifásico|Trifásico)\n+([A-Z\s,.-]+)\n+Leitura anterior', texto, re.DOTALL | re.IGNORECASE)
        if nome_match:
            dados_extraidos['Nome_Razao_Social'] = nome_match.group(1).strip()

        endereco_rua_match = re.search(r'Proxima Leitura\n+([^\n]+) DATAS DE', texto, re.DOTALL)
        if endereco_rua_match:
            dados_extraidos['Endereco_Rua_Numero'] = endereco_rua_match.group(1).strip()

        interior_line_match = re.search(r'LEITURAS.*?\n\s*(INTERIOR / ([A-Za-zÀ-ÖØ-öø-ÿ\s,.-]+))-([A-Z]{2})', texto, re.DOTALL)
        if interior_line_match:
            dados_extraidos['Bairro'] = 'INTERIOR'
            dados_extraidos['Cidade'] = interior_line_match.group(2).strip()
            dados_extraidos['Estado'] = interior_line_match.group(3).strip()

        match_cpf_cnpj = re.search(r'CPF/CNPJ:\s*([\d*]{3}\.[\d*]{3}\.[\d*]{3}-\d{2}|\d{2}\.[\d*]{3}\.[\d*]{3}\/\d{4}-\d{2})', texto)
        if match_cpf_cnpj:
            dados_extraidos['CNPJ_CPF'] = match_cpf_cnpj.group(1)

        match_cep = re.search(r'CEP:\s*(\d{2}\s*\d{3}-\d{3})', texto)
        if match_cep:
            dados_extraidos['CEP'] = match_cep.group(1)

        uc_match = re.search(r'UNIDADE CONSUMIDORA\n+Rota:\s*\d+,\s*Sequência:\s*\d+\s*([\d-]+)', texto, re.DOTALL)
        if uc_match:
            dados_extraidos['UC'] = uc_match.group(1).strip()

    except Exception as e:
        print(f"Erro ao extrair dados do sub-layout 'Cooperluz (sem COD UA)': {e}")
        pass
    return dados_extraidos


def _extrair_dados_layout_cooperluz_style(texto):
    if re.search(r'COD UA \d+', texto):
        return _extrair_dados_layout_cooperluz_sublayout_com_cod_ua(texto)
    else:
        return _extrair_dados_layout_cooperluz_sublayout_sem_cod_ua(texto)

def _extrair_dados_layout_coop_similar_style(texto, distributor_name):
    dados_extraidos = {
        'Nome_Razao_Social': 'Não encontrado', 'Endereco_Rua_Numero': 'Não encontrado',
        'Bairro': 'Não encontrado', 'CNPJ_CPF': 'Não encontrado',
        'Cidade': 'Não encontrado', 'CEP': 'Não encontrado',
        'UC': 'Não encontrado', 'Grupo_Tarifario': 'Não encontrado',
        'Tensao_Nominal_V': 'Não encontrado', 'Classe_Tarifaria': 'Não encontrado',
        'Estado': 'Não encontrado',
    }

    try:
        tipo_fornecimento_match = re.search(r'Tipo de Fornecimento:\s*(?:[\s\S]*?)(Monofásico|Bifásico|Trifásico)', texto, re.IGNORECASE | re.DOTALL)
        if tipo_fornecimento_match:
            tipo_fornecimento_extraido = tipo_fornecimento_match.group(1).strip()
            if 'Bifásico' in tipo_fornecimento_extraido or 'Monofásico' in tipo_fornecimento_extraido:
                dados_extraidos['Tensao_Nominal_V'] = 220
            elif 'Trifásico' in tipo_fornecimento_extraido:
                dados_extraidos['Tensao_Nominal_V'] = 380

        classificacao_line_match = re.search(r'Classificaç(?:ão|ao):\s*(.*?)(?:(?=\nTipo de Fornecimento)|\n|$)', texto, re.DOTALL | re.IGNORECASE)
        if classificacao_line_match:
            classif_line_content = classificacao_line_match.group(1).strip()
            match_grupo = re.search(r'(B[1-4]|A)', classif_line_content)
            if match_grupo:
                dados_extraidos['Grupo_Tarifario'] = match_grupo.group(1)
            match_classe = re.search(r'(Residencial|Comercial|Industrial|Rural|Poder Público|Iluminação Pública)', classif_line_content, re.IGNORECASE)
            if match_classe:
                dados_extraidos['Classe_Tarifaria'] = match_classe.group(1)

        nome_match = re.search(r'(?:Monofásico|Bifásico|Trifásico)\n+([A-Z\s,.-]+)\n+(?:Leitura anterior|DATAS DE)', texto, re.DOTALL | re.IGNORECASE)
        if nome_match:
            dados_extraidos['Nome_Razao_Social'] = nome_match.group(1).strip()

        endereco_rua_match = re.search(r'Proxima Leitura\n+([^\n]+) DATAS DE', texto, re.DOTALL)
        if endereco_rua_match:
            dados_extraidos['Endereco_Rua_Numero'] = endereco_rua_match.group(1).strip()

        interior_line_match = re.search(r'(?:LEITURAS|UNIDADE CONSUMIDORA).*?\n\s*(RURAL|INTERIOR)\s*/\s*([A-Za-zÀ-ÖØ-öø-ÿ\s,.-]+)-([A-Z]{2})', texto, re.DOTALL)
        if interior_line_match:
            dados_extraidos['Bairro'] = interior_line_match.group(1).strip()
            dados_extraidos['Cidade'] = interior_line_match.group(2).strip()
            dados_extraidos['Estado'] = interior_line_match.group(3).strip()

        match_cpf_cnpj = re.search(r'CPF/CNPJ:\s*([\d*]{3}\.[\d*]{3}\.[\d*]{3}-\d{2}|\d{2}\.[\d*]{3}\.[\d*]{3}\/\d{4}-\d{2})', texto)
        if match_cpf_cnpj:
            dados_extraidos['CNPJ_CPF'] = match_cpf_cnpj.group(1)

        match_cep = re.search(r'CEP:\s*(\d{2}\s*\d{3}-\d{3})', texto)
        if match_cep:
            dados_extraidos['CEP'] = match_cep.group(1)

        uc_match_explicit = re.search(r'UC:\s*([\d]+)[- ]', texto)
        if uc_match_explicit:
            dados_extraidos['UC'] = uc_match_explicit.group(1).strip()
        else:
            uc_match_rota = re.search(r'UNIDADE CONSUMIDORA\n+Rota:\s*\d+,\s*Sequência:\s*\d+\s*([\d]+)', texto, re.DOTALL)
            if uc_match_rota:
                dados_extraidos['UC'] = uc_match_rota.group(1).strip()
            else:
                uc_match_codigo_cliente = re.search(r'CÓDIGO DO CLIENTE\n*([\d]+)', texto)
                if uc_match_codigo_cliente:
                    dados_extraidos['UC'] = uc_match_codigo_cliente.group(1).strip()

    except Exception as e:
        print(f"Erro ao extrair dados do sub-layout '{distributor_name} (similar Cooperluz)': {e}")
        pass
    return dados_extraidos


def extrair_dados_fatura(caminho_pdf, distributor_type):
    try:
        with pdfplumber.open(caminho_pdf) as pdf:
            texto = pdf.pages[0].extract_text()

            if distributor_type == 'RGE':
                is_layout_adriano_style = re.search(r'Inscrição no CNPJ: \d{2}\.\d{3}\.\d{3}\/\d{4}-\d{2}\n+([A-Z\s,.]+)\n.*?Pelo CPF:\s*\d{3}\.\d{3}\.\d{3}-\d{2}', texto, re.DOTALL)
                is_layout_adroaldo_aire_style = re.search(r'Inscrição no CNPJ: \d{2}\.\d{3}\.\d{3}\/\d{4}-\d{2}\n+([A-Z\s,.]+)\n.*?CPF:\s*(((\*)?){6}\.\d{3}-(((\*)?){2}))', texto, re.DOTALL)
                is_layout_arcindo_style = re.search(r'DANFE - DOCUMENTO AUXILIAR DA NOTA FISCAL ELETRÔNICA', texto) and re.search(r'CÓDIGO DA UNIDADE CONSUMIDORA:', texto)

                if is_layout_adriano_style:
                    return _extrair_dados_layout_adriano_style(texto)
                elif is_layout_adroaldo_aire_style:
                    return _extrair_dados_layout_adroaldo_style(texto)
                elif is_layout_arcindo_style:
                    return _extrair_dados_layout_arcindo_style(texto)
                else:
                    extraction_functions = [
                        (_extrair_dados_layout_adriano_style, 'Adriano Style Fallback'),
                        (_extrair_dados_layout_adroaldo_style, 'Adroaldo/Aire Style Fallback'),
                        (_extrair_dados_layout_arcindo_style, 'Arcindo Style Fallback')
                    ]
                    best_match_data = None
                    max_found_fields = 0
                    for func, _ in extraction_functions:
                        current_data = func(texto)
                        current_found_fields = sum(1 for k, v in current_data.items() if v != 'Não encontrado' and k not in ['error'])
                        if current_found_fields > max_found_fields:
                            max_found_fields = current_found_fields
                            best_match_data = current_data
                    if best_match_data and max_found_fields > 0:
                         return best_match_data
                    
                    return {'error': f"Não foi possível identificar o layout da fatura RGE '{os.path.basename(caminho_pdf)}'. Layout desconhecido ou estrutura muito diferente."}

            elif distributor_type == 'COOPERLUZ':
                dados = _extrair_dados_layout_cooperluz_style(texto)
                if dados.get('Nome_Razao_Social') == 'Não encontrado':
                    return {'error': f"A fatura da Cooperluz '{os.path.basename(caminho_pdf)}' não pôde ser extraída. Nome/Razão Social não encontrado."}
                return dados
            elif distributor_type == 'CERTHIL':
                dados = _extrair_dados_layout_coop_similar_style(texto, 'CERTHIL')
                if dados.get('Nome_Razao_Social') == 'Não encontrado':
                     return {'error': f"A fatura da Certhil '{os.path.basename(caminho_pdf)}' não pôde ser extraída. Nome/Razão Social não encontrado."}
                return dados
            elif distributor_type == 'CERMISSOES':
                dados = _extrair_dados_layout_coop_similar_style(texto, 'CERMISSOES')
                if dados.get('Nome_Razao_Social') == 'Não encontrado':
                     return {'error': f"A fatura da Cermissões '{os.path.basename(caminho_pdf)}' não pôde ser extraída. Nome/Razão Social não encontrado."}
                return dados
            else:
                return {'error': f"Tipo de distribuidora '{distributor_type}' desconhecido."}

    except pdfplumber.pdfminer.pdfdocument.PDFSyntaxError:
        return {'error': f"O arquivo '{os.path.basename(caminho_pdf)}' não é um PDF válido ou está corrompido."}
    except FileNotFoundError:
        return {'error': f"O arquivo '{os.path.basename(caminho_pdf)}' não foi encontrado."}
    except Exception as e:
        return {'error': f"Erro inesperado durante a leitura do PDF: {e}"}

# --- Função auxiliar para parsear endereço para o Excel ---
def parse_address_for_excel(full_address):
    street = full_address
    number = ''
    match_number = re.search(r'(,\s*|\s*)(\d+[A-Za-z]?)\s*$', full_address)
    if match_number:
        number = match_number.group(2).strip()
        street = full_address[:match_number.start(1)].strip()
    elif "S/N" in full_address.upper():
        number = "S/N"
        street = full_address.upper().replace("S/N", "").replace(",","").strip()
    return street, number

# --- Função para converter graus decimais para GMS ---
def decimal_to_dms(decimal_degrees, is_latitude=True):
    if not isinstance(decimal_degrees, (int, float)):
        try:
            decimal_degrees = float(str(decimal_degrees).replace(',', '.'))
        except (ValueError, TypeError):
            return "N/A"

    direction = ""
    if is_latitude:
        direction = "N" if decimal_degrees >= 0 else "S"
    else: # Longitude
        direction = "E" if decimal_degrees >= 0 else "W"

    abs_degrees = abs(decimal_degrees)
    degrees = int(abs_degrees)
    minutes = int((abs_degrees - degrees) * 60)
    seconds = ((abs_degrees - degrees) * 60 - minutes) * 60
    seconds = round(seconds, 2)

    return f"""{degrees}° {minutes}' {seconds:.2f}" {direction}"""

# --- TABELA DE LOOKUP PARA CATEGORIA DA ENTRADA ELÉTRICA ---
TABELA_CATEGORIA_ELET = {
    'RIC BT - A2': {'Fases': 'Monofásico', 'Ramal de entrada': '10mm²', 'Disjuntor': '40A'},
    'GED 13 - A3': {'Fases': 'Monofásico', 'Ramal de entrada': '6mm²', 'Disjuntor': '32A'},
    'GED 13 - A4': {'Fases': 'Monofásico', 'Ramal de entrada': '16mm²', 'Disjuntor': '63A'},
    'RIC BT - B2': {'Fases': 'Bifásico', 'Ramal de entrada': '10mm²', 'Disjuntor': '50A'},
    'GED 13 - B3': {'Fases': 'Bifásico', 'Ramal de entrada': '10mm²', 'Disjuntor': '40A'},
    'RIC BT - B4': {'Fases': 'Bifásico', 'Ramal de entrada': '16mm²', 'Disjuntor': '50A'},
    'RIC BT - B5': {'Fases': 'Bifásico', 'Ramal de entrada': '25mm²', 'Disjuntor': '70A'},
    'GED 13 - C7': {'Fases': 'Trifásico', 'Ramal de entrada': '10mm²', 'Disjuntor': '40A'},
    'GED 13 - C8': {'Fases': 'Trifásico', 'Ramal de entrada': '16mm²', 'Disjuntor': '63A'},
    'GED 13 - C9': {'Fases': 'Trifásico', 'Ramal de entrada': '25mm²', 'Disjuntor': '80A'},
    'GED 13 - C10': {'Fases': 'Trifásico', 'Ramal de entrada': '35mm²', 'Disjuntor': '100A'},
    'GED 13 - C11': {'Fases': 'Trifásico', 'Ramal de entrada': '50mm²', 'Disjuntor': '125A'},
    'RIC BT - C13': {'Fases': 'Trifásico', 'Ramal de entrada': '10mm²', 'Disjuntor': '30A'},
    'RIC BT - C14': {'Fases': 'Trifásico', 'Ramal de entrada': '10mm²', 'Disjuntor': '40A'},
    'RIC BT - C15': {'Fases': 'Trifásico', 'Ramal de entrada': '16mm²', 'Disjuntor': '50A'},
    'RIC BT - C16': {'Fases': 'Trifásico', 'Ramal de entrada': '25mm²', 'Disjuntor': '70A'},
    'RIC BT - C17': {'Fases': 'Trifásico', 'Ramal de entrada': '35mm²', 'Disjuntor': '100A'},
    'RIC BT - C18': {'Fases': 'Trifásico', 'Ramal de entrada': '50mm²', 'Disjuntor': '125A'},
}

# --- TABELA DE LOOKUP PARA PARÂMETROS AC DO INVERSOR (BASEADO NA POTÊNCIA) ---
TABELA_PARAMETROS_INVERSOR_AC = [
    {'min_power': 0.0, 'max_power': 3.9, 'DISJUNTOR_CA': '16A', 'DISJ_CA_ATEN': '80A a 160A', 'DISJ_CA_INTR': '3kA', 'DISJ_CA_TENS': '400V', 'CABO_CA': '4mm²'},
    {'min_power': 4.0, 'max_power': 6.9, 'DISJUNTOR_CA': '32A', 'DISJ_CA_ATEN': '160A a 320A', 'DISJ_CA_INTR': '3kA', 'DISJ_CA_TENS': '400V', 'CABO_CA': '6mm²'},
    {'min_power': 7.0, 'max_power': 8.9, 'DISJUNTOR_CA': '40A', 'DISJ_CA_ATEN': '200A a 400A', 'DISJ_CA_INTR': '3kA', 'DISJ_CA_TENS': '400V', 'CABO_CA': '10mm²'},
    {'min_power': 9.0, 'max_power': 10.9, 'DISJUNTOR_CA': '50A', 'DISJ_CA_ATEN': '250A a 500A', 'DISJ_CA_INTR': '3kA', 'DISJ_CA_TENS': '400V', 'CABO_CA': '16mm²'},
    {'min_power': 12.0, 'max_power': 29.9, 'DISJUNTOR_CA': '32A', 'DISJ_CA_ATEN': '160A a 320A', 'DISJ_CA_INTR': '3kA', 'DISJ_CA_TENS': '400V', 'CABO_CA': '6mm²'},
    {'min_power': 30.0, 'max_power': 39.9, 'DISJUNTOR_CA': '63A', 'DISJ_CA_ATEN': '315A a 630A', 'DISJ_CA_INTR': '3kA', 'DISJ_CA_TENS': '400V', 'CABO_CA': '16mm²'},
    {'min_power': 40.0, 'max_power': 49.9, 'DISJUNTOR_CA': '70A', 'DISJ_CA_ATEN': '350A a 700A', 'DISJ_CA_INTR': '3kA', 'DISJ_CA_TENS': '400V', 'CABO_CA': '25mm²'},
    {'min_power': 50.0, 'max_power': 59.9, 'DISJUNTOR_CA': '100A', 'DISJ_CA_ATEN': '500A a 1000A', 'DISJ_CA_INTR': '6kA', 'DISJ_CA_TENS': '400V', 'CABO_CA': '35mm²'},
    {'min_power': 60.0, 'max_power': 75.0, 'DISJUNTOR_CA': 'Verificação pontual', 'DISJ_CA_ATEN': 'Verificação pontual', 'DISJ_CA_INTR': 'Verificação pontual', 'DISJ_CA_TENS': 'Verificação pontual', 'CABO_CA': 'Verificação pontual'},
]

def get_ac_parameters_by_inverter_power(potencia_inversor_kw):
    if not isinstance(potencia_inversor_kw, (int, float)):
        return {
            'CABO_CA': 'Não calculado', 'DISJUNTOR_CA': 'Não calculado',
            'DISJ_CA_TENS': 'Não calculado', 'DISJ_CA_INTR': 'Não calculado',
            'DISJ_CA_ATEN': 'Não calculado'
        }

    for linha in TABELA_PARAMETROS_INVERSOR_AC:
        if linha['min_power'] <= potencia_inversor_kw <= linha['max_power']:
            return {
                'CABO_CA': linha['CABO_CA'], 'DISJUNTOR_CA': linha['DISJUNTOR_CA'],
                'DISJ_CA_TENS': linha['DISJ_CA_TENS'], 'DISJ_CA_INTR': linha['DISJ_CA_INTR'],
                'DISJ_CA_ATEN': linha['DISJ_CA_ATEN']
            }
    
    return {
        'CABO_CA': 'Fora da faixa (0-75kW)', 'DISJUNTOR_CA': 'Fora da faixa (0-75kW)',
        'DISJ_CA_TENS': 'Fora da faixa (0-75kW)', 'DISJ_CA_INTR': 'Fora da faixa (0-75kW)',
        'DISJ_CA_ATEN': 'Fora da faixa (0-75kW)'
    }

# --- FUNÇÃO PARA CALCULAR VARIÁVEIS DO SISTEMA EM PYTHON ---
def calculate_system_variables(all_data):
    calculated_data = {}

    categoria = all_data.get('CATEGORIA')
    if categoria and categoria in TABELA_CATEGORIA_ELET:
        calculated_data['NUMERO_FASES_CALCULADO'] = TABELA_CATEGORIA_ELET[categoria]['Fases']
        calculated_data['RAMAL_ENTRADA_CALCULADO'] = TABELA_CATEGORIA_ELET[categoria]['Ramal de entrada']
        calculated_data['DISJUNTOR_CALCULADO'] = TABELA_CATEGORIA_ELET[categoria]['Disjuntor']
    else:
        calculated_data['NUMERO_FASES_CALCULADO'] = 'Não calculado'
        calculated_data['RAMAL_ENTRADA_CALCULADO'] = 'Não calculado'
        calculated_data['DISJUNTOR_CALCULADO'] = 'Não calculado'

    try:
        quantidade_placas = float(str(all_data.get('QUANTIDADE_PLACAS_MANUAL', 0)).replace(',', '.'))
        potencia_modulo_wp = float(str(all_data.get('POTENCIA_MODULO_MANUAL', 0)).replace(',', '.').replace('Wp', '').strip())
        calculated_data['POTENCIA_PICO_MODULOS'] = (quantidade_placas * potencia_modulo_wp) / 1000
    except (ValueError, TypeError):
        calculated_data['POTENCIA_PICO_MODULOS'] = 'Não calculado'

    potencia_inversor_kw_value = None
    try:
        quantidade_inversor = float(str(all_data.get('QUANTIDADE_INVERSOR_MANUAL', 0)).replace(',', '.'))
        potencia_inversor_kw_value = all_data.get('POTENCIA_INVERSOR_MANUAL')
        
        if isinstance(potencia_inversor_kw_value, (int, float)):
            calculated_data['POTENCIA_TOTAL_SISTEMA_KWP'] = quantidade_inversor * potencia_inversor_kw_value
        else:
            calculated_data['POTENCIA_TOTAL_SISTEMA_KWP'] = 'Não calculado'
            potencia_inversor_kw_value = 'Não calculado'
    except (ValueError, TypeError):
        calculated_data['POTENCIA_TOTAL_SISTEMA_KWP'] = 'Não calculado'
        potencia_inversor_kw_value = 'Não calculado'

    if isinstance(potencia_inversor_kw_value, (int, float)):
        ac_parameters = get_ac_parameters_by_inverter_power(potencia_inversor_kw_value)
        calculated_data.update(ac_parameters)
    else:
        calculated_data['CABO_CA'] = 'Não calculado'
        calculated_data['DISJUNTOR_CA'] = 'Não calculado'
        calculated_data['DISJ_CA_TENS'] = 'Não calculado'
        calculated_data['DISJ_CA_INTR'] = 'Não calculado'
        calculated_data['DISJ_CA_ATEN'] = 'Não calculado'

    calculated_data['MODELO_MODULO_CALCULADO'] = all_data.get('MODELO_MODULO_CALCULADO', 'Não informado')
    calculated_data['AREA_ARRANJOS_CALCULADO'] = all_data.get('AREA_ARRANJOS_CALCULADO', 'Não informado')
    calculated_data['MODELO_INVERSOR_CALCULADO'] = all_data.get('MODELO_INVERSOR_CALCULADO', 'Não informado')
    calculated_data['INMETRO'] = all_data.get('INMETRO', 'Não informado')
    calculated_data['ISOLACAO_CA'] = all_data.get('ISOLACAO_CA', 'Não informado')

    return calculated_data

# --- Mapeamento de campos internos para células do Excel (Planilha Projetos FV) ---
EXCEL_PROJETO_FV_CELL_MAPPING = {
    'Nome_Razao_Social': 'B2', 'Endereco_Rua': 'C2', 'Numero_Endereco': 'D2', 'Bairro': 'E2',
    'E-MAIL': 'F2', 'TELEFONE': 'G2', 'CNPJ_CPF': 'H2', 'Cidade': 'I2', 'CEP': 'J2',
    'LATITUDE': 'K2', 'LONGITUDE': 'L2', 'UC': 'A5', 'Grupo_Tarifario': 'B5',
    'Classe_Tarifaria': 'C5', 'Tensao_Nominal_V': 'D5', 'CARGA_INSTALADA': 'E5',
    'CATEGORIA': 'F5', 'TIPO_DE_ATENDIMENTO': 'G5', 'TIPO_DE_CAIXA': 'A8',
    'ISOLACAO': 'I8', 'POTENCIA_MODULO_MANUAL': 'A27', 'FABRICANTE_MODULO_MANUAL': 'C27',
    'QUANTIDADE_PLACAS_MANUAL': 'D27', 'MODELO_MODULO_CALCULADO': 'B27',
    'AREA_ARRANJOS_CALCULADO': 'G27', 'POTENCIA_INVERSOR_MANUAL': 'A29',
    'FABRICANTE_INVERSOR_MANUAL': 'C29', 'QUANTIDADE_INVERSOR_MANUAL': 'D29',
    'MODELO_INVERSOR_CALCULADO': 'B29', 'INMETRO': 'J29', 'NUMERO_FASES_CALCULADO': 'B8',
    'RAMAL_ENTRADA_CALCULADO': 'D8', 'DISJUNTOR_CALCULADO': 'E11',
    'POTENCIA_PICO_MODULOS': 'F27', 'POTENCIA_TOTAL_SISTEMA_KWP': 'G29',
    'CABO_CA': 'F44', 'DISJUNTOR_CA': 'A44', 'DISJ_CA_TENS': 'C44',
    'DISJ_CA_INTR': 'I44', 'DISJ_CA_ATEN': 'J44', 'ISOLACAO_CA': 'H44',
}

# --- Mapeamento dos textos (placeholders) no Anexo F para as variáveis do Python ---
ANEXO_F_PLACEHOLDER_TO_PYTHON_VAR = {
    'NOME/RAZÃO SOCIAL': 'Nome_Razao_Social', 'CNPJ/CPF': 'CNPJ_CPF', 'UC': 'UC',
    'ENDEREÇO + N° + BAIRRO': 'FULL_ADDRESS_COMPOSED', 'CEP': 'CEP', 'CIDADE': 'Cidade',
    'TELEFONE': 'TELEFONE', 'E-MAIL': 'E-MAIL', 'CATEGORIA': 'CATEGORIA',
    'TIPO DE ATENDIMENTO': 'TIPO_DE_ATENDIMENTO', 'TIPO DE CAIXA': 'TIPO_DE_CAIXA',
    'CARGA INSTALADA': 'CARGA_INSTALADA', 'DISJUNTOR SIMPL.': 'DISJUNTOR_CALCULADO',
    'QUANTIDADE PLACAS': 'QUANTIDADE_PLACAS_MANUAL', 'FABRICANTE_MODULO': 'FABRICANTE_MODULO_MANUAL',
    'MODELO_MODULO': 'MODELO_MODULO_CALCULADO', 'ÁREA': 'AREA_ARRANJOS_CALCULADO',
    'QUANTIDADE INVERSOR': 'QUANTIDADE_INVERSOR_MANUAL', 'FABRICANTE_INVERSOR': 'FABRICANTE_INVERSOR_MANUAL',
    'MODELO_INVERSOR': 'MODELO_INVERSOR_CALCULADO', 'POTÊNCIA (ANEXO I)': 'POTENCIA_TOTAL_SISTEMA_KWP',
    'POTÊNCIA NOMINAL': 'POTENCIA_INVERSOR_MANUAL', 'DATA OPERAÇÃO': 'DATA_OPERACAO_PREVISTA',
    'N° DE FASES': 'NUMERO_FASES_CALCULADO', 'RAMAL DE ENTRADA': 'RAMAL_ENTRADA_CALCULADO',
    'POTENCIA_PICO_MODULOS': 'POTENCIA_PICO_MODULOS', 'LATITUDE': 'LATITUDE',
    'LONGITUDE': 'LONGITUDE', 'LATITUDE_GMS': 'LATITUDE_GMS', 'LONGITUDE_GMS': 'LONGITUDE_GMS',
    'INMETRO': 'INMETRO',
}

# Mapeamento dos textos (placeholders) no Anexo I para as variáveis do Python ---
ANEXO_I_PLACEHOLDER_TO_PYTHON_VAR = {
    'UC': 'UC', 'GRUPO_TARIFARIO': 'Grupo_Tarifario', 'CLASSE_TARIFARIA': 'Classe_Tarifaria',
    'ENDERECO_RUA_NUMERO': 'Endereco_Rua_Numero', 'ENDERECO_RUA_NUMERO_BAIRRO': 'Endereco_Rua_Numero_Bairro',
    'BAIRRO': 'Bairro', 'CIDADE': 'Cidade', 'ESTADO': 'Estado', 'CIDADE_ESTADO': 'CIDADE_ESTADO',
    'CEP': 'CEP', 'LATITUDE_GMS': 'LATITUDE_GMS', 'LONGITUDE_GMS': 'LONGITUDE_GMS',
    'NOME_RAZAO_SOCIAL': 'Nome_Razao_Social', 'CNPJ_CPF': 'CNPJ_CPF', 'TELEFONE': 'TELEFONE',
    'EMAIL': 'E-MAIL', 'CARGA_INSTALADA': 'CARGA_INSTALADA', 'TENSAO_NOMINAL_V': 'Tensao_Nominal_V',
    'TIPO_DE_ATENDIMENTO': 'TIPO_DE_ATENDIMENTO', 'DISJUNTOR_CALCULADO': 'DISJUNTOR_CALCULADO',
    'RAMAL_ENTRADA_CALCULADO': 'RAMAL_ENTRADA_CALCULADO',
    'QUANTIDADE_PLACAS_MANUAL': 'QUANTIDADE_PLACAS_MANUAL', 'POTENCIA_MODULO_MANUAL': 'POTENCIA_MODULO_MANUAL',
    'POTENCIA_PICO_MODULOS': 'POTENCIA_PICO_MODULOS', 'FABRICANTE_MODULO_MANUAL': 'FABRICANTE_MODULO_MANUAL',
    'MODELO_MODULO_CALCULADO': 'MODELO_MODULO_CALCULADO',
    'QUANTIDADE_INVERSOR_MANUAL': 'QUANTIDADE_INVERSOR_MANUAL', 'POTENCIA_INVERSOR_MANUAL': 'POTENCIA_INVERSOR_MANUAL',
    'POTENCIA_TOTAL_SISTEMA_KWP': 'POTENCIA_TOTAL_SISTEMA_KWP',
    'FABRICANTE_INVERSOR_MANUAL': 'FABRICANTE_INVERSOR_MANUAL',
    'MODELO_INVERSOR_CALCULADO': 'MODELO_INVERSOR_CALCULADO', 'AREA_ARRANJOS_CALCULADO': 'AREA_ARRANJOS_CALCULADO',
    'ART': 'ART', 'DATA_ATUAL': 'DATA_ATUAL', 'DATA_OPERACAO_PREVISTA': 'DATA_OPERACAO_PREVISTA',
    'POTENCIA_MODULO_MANUAL_KWP': 'POTENCIA_MODULO_MANUAL_KWP',
    'NUMERO_FASES_CALCULADO': 'NUMERO_FASES_CALCULADO', 'INMETRO': 'INMETRO', 'ISOLACAO_CA': 'ISOLACAO_CA',
    'CABO_CA': 'CABO_CA', 'DISJUNTOR_CA': 'DISJUNTOR_CA', 'DISJ_CA_TENS': 'DISJ_CA_TENS',
    'DISJ_CA_INTR': 'DISJ_CA_INTR', 'DISJ_CA_ATEN': 'DISJ_CA_ATEN',
}


# --- Configuração do Flask ---
app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.secret_key = 'supersecretkey'

ALLOWED_EXTENSIONS = {'pdf'}
ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg'}

EXCEL_PROJETO_FV_TEMPLATE_FILENAME = 'Planilha Projetos FV.xlsx'
EXCEL_PROJETO_FV_TEMPLATE_PATH = os.path.join(app.root_path, 'templates', EXCEL_PROJETO_FV_TEMPLATE_FILENAME)

ANEXO_F_TEMPLATE_FILENAME = 'Formulário Anexo F - GED 15303.xlsx'
ANEXO_F_TEMPLATE_PATH = os.path.join(app.root_path, 'templates', ANEXO_F_TEMPLATE_FILENAME)
ANEXO_F_SHEET_NAME = 'Anexo F - GED 15303'

ANEXO_I_TEMPLATE_FILENAME = 'Anexo I.xlsx'
ANEXO_I_TEMPLATE_PATH = os.path.join(app.root_path, 'templates', ANEXO_I_TEMPLATE_FILENAME)
ANEXO_I_SHEET_NAME = 'Plan1'

ANEXO_E_TEMPLATE_FILENAME = 'Formulário Anexo E - GED.docx'
ANEXO_E_TEMPLATE_PATH = os.path.join(app.root_path, 'templates', ANEXO_E_TEMPLATE_FILENAME)

TERMO_ACEITE_TEMPLATE_FILENAME = 'Termo de Aceite.docx'
TERMO_ACEITE_TEMPLATE_PATH = os.path.join(app.root_path, 'templates', TERMO_ACEITE_TEMPLATE_FILENAME)

PROCURACAO_TEMPLATE_FILENAME = 'Procuracao.docx'
PROCURACAO_TEMPLATE_PATH = os.path.join(app.root_path, 'templates', PROCURACAO_TEMPLATE_FILENAME)

TERMO_ACEITE_INCISO_III_TEMPLATE_FILENAME = 'Termo de Aceite Inciso III.docx'
TERMO_ACEITE_INCISO_III_TEMPLATE_PATH = os.path.join(app.root_path, 'templates', TERMO_ACEITE_INCISO_III_TEMPLATE_FILENAME)

RESPONSABILIDADE_TECNICA_TEMPLATE_FILENAME = 'Responsabilidade Tecnica.docx'
RESPONSABILIDADE_TECNICA_TEMPLATE_PATH = os.path.join(app.root_path, 'templates', RESPONSABILIDADE_TECNICA_TEMPLATE_FILENAME)

DADOS_GD_UFV_TEMPLATE_FILENAME = 'Dados para GD de UFV.docx'
DADOS_GD_UFV_TEMPLATE_PATH = os.path.join(app.root_path, 'templates', DADOS_GD_UFV_TEMPLATE_FILENAME)

MEMORIAL_DESCRITIVO_TEMPLATE_FILENAME = 'Memorial Descritivo Padrão Fecoergs.docx'
MEMORIAL_DESCRITIVO_TEMPLATE_PATH = os.path.join(app.root_path, 'templates', MEMORIAL_DESCRITIVO_TEMPLATE_FILENAME)

CERTIDAO_REGISTRO_PROFISSIONAL_FILENAME = 'Certidao de Registro Profissional.pdf'
CERTIDAO_REGISTRO_PROFISSIONAL_PATH = os.path.join(app.root_path, 'templates', CERTIDAO_REGISTRO_PROFISSIONAL_FILENAME)
if not os.path.exists(CERTIDAO_REGISTRO_PROFISSIONAL_PATH):
    with open(CERTIDAO_REGISTRO_PROFISSIONAL_PATH, 'wb') as f:
        f.write(b'%PDF-1.4\n%\xc2\xa5\xc2\xb1\xc2\xae\xc2\xbb\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n2 0 obj<</Type/Pages/Count 0>>endobj\nxref\n0 3\n0000000000 65535 f\n0000000009 00000 n\n0000000074 00000 n\ntrailer<</Size 3/Root 1 0 R>>startxref\n120\n%%EOF')
    print(f"ATENÇÃO: Arquivo de template '{CERTIDAO_REGISTRO_PROFISSIONAL_FILENAME}' não encontrado na pasta 'templates'. Um placeholder PDF vazio foi criado.")

def allowed_file(filename, allowed_extensions):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

# --- FUNÇÃO PARA SUBSTITUIR PLACEHOLDERS NO DOCX COM PRESERVAÇÃO DE FORMATO ---
def replace_docx_placeholders(doc_path, replacements):
    document = Document(doc_path)
    bold_keys = ['NOME_RAZAO_SOCIAL', 'CPF_CNPJ', 'UC']

    def process_text_block(text_block_container):
        for p in text_block_container.paragraphs:
            if not p.text.strip():
                continue

            original_full_text = p.text
            
            default_run_format = {
                'bold': None, 'italic': None, 'font_name': None, 'font_size': None, 'underline': None
            }
            if p.runs:
                first_run = p.runs[0]
                default_run_format['bold'] = first_run.bold
                default_run_format['italic'] = first_run.italic
                default_run_format['font_name'] = first_run.font.name
                default_run_format['font_size'] = first_run.font.size
                default_run_format['underline'] = first_run.underline

            final_segments_for_runs = []
            
            placeholder_patterns = [re.escape('{{' + key + '}}') for key in replacements.keys()]
            combined_pattern = '|'.join(placeholder_patterns)
            
            last_idx = 0
            if combined_pattern:
                for match in re.finditer(combined_pattern, original_full_text):
                    if match.start() > last_idx:
                        final_segments_for_runs.append((original_full_text[last_idx:match.start()], False, None))
                    
                    matched_placeholder = match.group(0)
                    placeholder_key = matched_placeholder.strip('{}')
                    
                    value_to_insert = replacements.get(placeholder_key, matched_placeholder)
                    final_segments_for_runs.append((str(value_to_insert), True, placeholder_key))
                    
                    last_idx = match.end()
            
            if last_idx < len(original_full_text):
                final_segments_for_runs.append((original_full_text[last_idx:], False, None))

            if not (final_segments_for_runs and any(s[1] for s in final_segments_for_runs)) and not combined_pattern:
                return 

            for i in range(len(p.runs) - 1, -1, -1):
                p.runs[i]._element.getparent().remove(p.runs[i]._element)

            for text, is_replaced, placeholder_key in final_segments_for_runs:
                if text:
                    new_run = p.add_run(text)
                    new_run.bold = default_run_format['bold']
                    new_run.italic = default_run_format['italic']
                    new_run.font.name = default_run_format['font_name']
                    new_run.font.size = default_run_format['font_size']
                    new_run.underline = default_run_format['underline']

                    if is_replaced and placeholder_key in bold_keys:
                        new_run.bold = True
    
    process_text_block(document)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                process_text_block(cell)

    return document

# --- FUNÇÃO PARA GERAR CONTEÚDO DO TXT DA ART ---
def generate_art_txt_content(data):
    cnpj_cpf_numeric = re.sub(r'[^0-9]', '', str(data.get('CNPJ_CPF', '')))
    cep_numeric = re.sub(r'[^0-9]', '', str(data.get('CEP', '')))
    telefone_numeric = re.sub(r'[^0-9]', '', str(data.get('TELEFONE', '')))
    potencia_total_kwp_clean = format_value_for_display(data.get('POTENCIA_TOTAL_SISTEMA_KWP', '0'), is_numeric=True)

    return f"""Criar nova ART:
- No site do CREA acessar: ART WEB vá até: NOVA ART

Em nova ART:
-Professional:
Empresa executante da obra/serviço. Selecione: QUASAT SOLAR- EQUIPAMENTOS DE ENERGIA

-ART:
Tipo de ART: Obra ou Serviço
Motivo da ART: Normal

- Contratante:
Busque o contratante (flecha amarela)
Clique em CADASTRAR
Contratante: {data.get('Nome_Razao_Social', 'Não informado')}
CPF/CNPJ: {cnpj_cpf_numeric}
E-mail: {data.get('E-MAIL', 'Não informado')}
CEP: {cep_numeric}
Telefone: {telefone_numeric}
CONFIRME

- Obra/Serviço:
Proprietário: clique em "Buscar dados do Contratante"
Finalidade: Outras Finalidades
Valor Contrato(R$): 2000
Honorários(R$): 200
Data Início: {format_value_for_display(data.get('DATA_ART'), is_date=True, date_separator='-')}
Data Previsão de Fim: {format_value_for_display(data.get('DATA_OPERACAO_PREVISTA'), is_date=True, date_separator='-')}

- Atividades
Ativ. Téc.(flecha amarela): Projeto e Execução
Qtd.: {potencia_total_kwp_clean}
Unidade(flecha amarela): Quilowatt

CONFIRMAR - FINALIZAR

LATITUDE: {format_value_for_display(data.get('LATITUDE'), is_numeric=True, numeric_decimal_separator='.')}
LONGITUDE: {format_value_for_display(data.get('LONGITUDE'), is_numeric=True, numeric_decimal_separator='.')}
LATITUDE_GMS: {data.get('LATITUDE_GMS', 'Não informado')}
LONGITUDE_GMS: {data.get('LONGITUDE_GMS', 'Não informado')}
"""

# --- FUNÇÃO PARA GERAR CONTEÚDO DO TXT DA POSTAGEM (GENÉRICA) ---
def generate_postagem_txt_content(data):
    potencia_total_kwp_clean = format_value_for_display(data.get('POTENCIA_TOTAL_SISTEMA_KWP', '0'), is_numeric=True)
    potencia_pico_modulos_clean = format_value_for_display(data.get('POTENCIA_PICO_MODULOS', '0'), is_numeric=True)
    potencia_inversor_manual_clean = format_value_for_display(data.get('POTENCIA_INVERSOR_MANUAL', '0'), is_numeric=True)
    
    return f"""Instalação: {data.get('UC', 'Não informado')}
Título do projeto: FV {data.get('Nome_Razao_Social', 'Não informado')}
Potência instalada do Gerador: {potencia_total_kwp_clean}
Carga atual do cliente: {format_value_for_display(data.get('CARGA_INSTALADA'), is_numeric=True)}
Latitude: {format_value_for_display(data.get('LATITUDE'), is_numeric=True, numeric_decimal_separator='.')}
Longitude: {format_value_for_display(data.get('LONGITUDE'), is_numeric=True, numeric_decimal_separator='.')}
Data prevista para ligação: {format_value_for_display(data.get('DATA_OPERACAO_PREVISTA'), is_date=True)}

Tipo de solicitação: SOLICITAÇÃO DE ACESSO
Quantidade de Beneficiárias: 0
Telefone do Titular (DDD + Número): {data.get('TELEFONE', 'Não informado')}
E-mail do Titular: {data.get('E-MAIL', 'Não informado')}
Latitude: {format_value_for_display(data.get('LATITUDE'), is_numeric=True, numeric_decimal_separator='.')}
Longitude: {format_value_for_display(data.get('LONGITUDE'), is_numeric=True, numeric_decimal_separator='.')}
Potência Total dos Módulos (kw): {potencia_pico_modulos_clean}
Quantidade de Módulos: {format_value_for_display(data.get('QUANTIDADE_PLACAS_MANUAL'), is_numeric=True)}
Potência Total dos Inversores (kw): {potencia_inversor_manual_clean}
Quantidade de Inversores: {format_value_for_display(data.get('QUANTIDADE_INVERSOR_MANUAL'), is_numeric=True)}
Fabricante(s) dos Módulos: {format_value_for_display(data.get('FABRICANTE_MODULO_MANUAL'))}
Modelo(s) dos Módulos: {format_value_for_display(data.get('MODELO_MODULO_CALCULADO'))}
Área Total dos Arranjos (m²): {format_value_for_display(data.get('AREA_ARRANJOS_CALCULADO'), is_numeric=True)}
Fabricante(s) dos Inversores: {format_value_for_display(data.get('FABRICANTE_INVERSOR_MANUAL'))}
Modelo(s) dos Inversores: {format_value_for_display(data.get('MODELO_INVERSOR_CALCULADO'))}


E-mail: {data.get('E-MAIL', 'Não informado')}
Telefone: {data.get('TELEFONE', 'Não informado')}
Celular: {data.get('TELEFONE', 'Não informado')}
Documento de Responsabilidade Técnica:
Data do Documento: {format_value_for_display(data.get('DATA_ART'), is_date=True, date_separator='-')}
Email: projetos@quasatservices.com

LATITUDE_GMS: {data.get('LATITUDE_GMS', 'Não informado')}
LONGITUDE_GMS: {data.get('LONGITUDE_GMS', 'Não informado')}
"""

def generate_images_pdf(image_data_list, output_pdf_path):
    c = canvas.Canvas(output_pdf_path, pagesize=portrait(A4))
    width, height = portrait(A4)
    margin = 2 * cm

    for item in image_data_list:
        img_path = item['path']
        title = item['title']

        if not os.path.exists(img_path):
            c.setFont('Helvetica-Bold', 14)
            c.drawCentredString(width / 2.0, height - 2 * cm, title)
            c.setFont('Helvetica', 12)
            c.drawCentredString(width / 2.0, height / 2.0, f"Imagem não encontrada: {title}")
            c.showPage()
            continue

        try:
            c.setFont('Helvetica-Bold', 14)
            c.drawCentredString(width / 2.0, height - cm, title)

            pil_img = Image.open(img_path)
            max_img_width = width - 2 * margin
            max_img_height = height - 3.5 * cm

            img_width_orig, img_height_orig = pil_img.size
            
            if img_width_orig > max_img_width or img_height_orig > max_img_height:
                scale_width = max_img_width / img_width_orig
                scale_height = max_img_height / img_height_orig
                scale_factor = min(scale_width, scale_height)

                img_width = img_width_orig * scale_factor
                img_height = img_height_orig * scale_factor
            else:
                img_width = img_width_orig
                img_height = img_height_orig

            x_pos = (width - img_width) / 2
            y_pos = (height - img_height) / 2 - (1 * cm)

            c.drawImage(ImageReader(img_path), x_pos, y_pos, width=img_width, height=img_height, preserveAspectRatio=True)
            c.showPage()

        except Exception as e:
            print(f"Erro ao adicionar imagem {img_path} ao PDF: {e}")
            c.setFont('Helvetica-Bold', 14)
            c.drawCentredString(width / 2.0, height - 2 * cm, title)
            c.setFont('Helvetica', 12)
            c.drawCentredString(width / 2.0, height / 2.0, f"Erro ao processar imagem: {title}")
            c.showPage()
    c.save()


@app.route('/', methods=['GET'])
def upload_form():
    session.clear()
    return render_template('index.html')

@app.route('/process_data', methods=['GET'])
def show_process_data_form():
    dados = session.pop('all_input_data_for_correction', None) or session.pop('extracted_data_from_pdf', {})
    
    if not dados and not session.get('extracted_data_from_pdf'):
        return redirect(url_for('upload_form'))

    session['current_process_data_form_data'] = dados 
    
    return render_template('process_data.html', dados=dados, TABELA_CATEGORIA_ELET_KEYS=list(TABELA_CATEGORIA_ELET.keys()))

@app.route('/clear_session_and_redirect_to_upload')
def clear_session_and_redirect_to_upload():
    session.clear()
    return redirect(url_for('upload_form'))

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return render_template('index.html', error='Nenhum arquivo enviado.'), 400

    file = request.files['file']
    distributor_type = request.form.get('distribuidora')

    if not distributor_type:
        return render_template('index.html', error='Por favor, selecione a distribuidora.'), 400

    if file.filename == '':
        return render_template('index.html', error='Nenhum arquivo selecionado.'), 400

    if file and allowed_file(file.filename, ALLOWED_EXTENSIONS):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        dados_fatura = extrair_dados_fatura(filepath, distributor_type)

        if 'error' in dados_fatura or dados_fatura.get('Nome_Razao_Social') == 'Não encontrado':
            os.remove(filepath)
            error_message = dados_fatura.get('error', f'Erro desconhecido na extração da fatura para {distributor_type}. O Nome/Razão Social não foi encontrado, indicando um problema com o layout ou a legibilidade do PDF.')
            return render_template('index.html', error=error_message), 500

        os.remove(filepath)

        dados_fatura['distributor_type'] = distributor_type
        session['extracted_data_from_pdf'] = dados_fatura
        
        return redirect(url_for('show_process_data_form'))
    else:
        return render_template('index.html', error='Tipo de arquivo não permitido. Por favor, envie um PDF.'), 400

NUMERIC_KEYS_FOR_FORMATTING = {
    'LATITUDE', 'LONGITUDE', 'POTENCIA_MODULO_MANUAL', 'POTENCIA_INVERSOR_MANUAL',
    'CARGA_INSTALADA', 'POTENCIA_PICO_MODULOS', 'POTENCIA_TOTAL_SISTEMA_KWP',
    'POTENCIA_MODULO_MANUAL_KWP', 'AREA_ARRANJOS_CALCULADO', 'Tensao_Nominal_V',
    'DISJUNTOR_CALCULADO', 'DISJUNTOR_CA', 'DISJ_CA_TENS', 'DISJ_CA_INTR', 'DISJ_CA_ATEN'
}

DATE_KEYS_FOR_FORMATTING = {
    'DATA_ATUAL', 'DATA_OPERACAO_PREVISTA', 'DATA_ART'
}

def get_formatted_value_for_doc(key, data_dict):
    value = data_dict.get(key)
    if key in NUMERIC_KEYS_FOR_FORMATTING:
        return format_value_for_display(value, is_numeric=True, numeric_decimal_separator='.')
    elif key in DATE_KEYS_FOR_FORMATTING:
        return format_value_for_display(value, is_date=True)
    return str(value) if value is not None else ''


@app.route('/process_and_save', methods=['POST'])
def process_and_save():
    all_input_data = session.pop('current_process_data_form_data', {})
    if not all_input_data:
        return render_template('index.html', error='Dados da sessão expirados ou não encontrados. Por favor, reinicie o processo.'), 400

    distributor_type = request.form.get('distributor_type', all_input_data.get('distributor_type', 'RGE'))
    all_input_data['distributor_type'] = distributor_type

    for key, value in request.form.items():
        processed_value = value.strip() if isinstance(value, str) else value

        if key in ['CARGA_INSTALADA', 'POTENCIA_MODULO_MANUAL', 'POTENCIA_INVERSOR_MANUAL', 
                    'LATITUDE', 'LONGITUDE', 'AREA_ARRANJOS_CALCULADO']:
            if processed_value:
                try:
                    all_input_data[key] = float(str(processed_value).replace(',', '.'))
                except ValueError:
                    all_input_data[key] = processed_value
            else:
                all_input_data[key] = None
        elif key in ['QUANTIDADE_PLACAS_MANUAL', 'QUANTIDADE_INVERSOR_MANUAL', 'ART']:
            if processed_value and str(processed_value).isdigit():
                all_input_data[key] = int(processed_value)
            else:
                all_input_data[key] = processed_value
        elif key.startswith('extracted_'):
            original_key = key[len('extracted_'):]
            if original_key == 'Tensao_Nominal_V' and processed_value and processed_value.isdigit():
                all_input_data[original_key] = int(processed_value)
            else:
                all_input_data[original_key] = processed_value
        else:
            all_input_data[key] = processed_value
    
    calculated_values = calculate_system_variables(all_input_data)
    all_input_data.update(calculated_values)

    bairro = all_input_data.get('Bairro', '').strip()
    cidade = all_input_data.get('Cidade', '').strip()
    if bairro and cidade:
        all_input_data['BAIRRO_CIDADE_COMBINADO'] = f"{bairro} - {cidade}"
    elif bairro:
        all_input_data['BAIRRO_CIDADE_COMBINADO'] = bairro
    elif cidade:
        all_input_data['BAIRRO_CIDADE_COMBINADO'] = cidade
    else:
        all_input_data['BAIRRO_CIDADE_COMBINADO'] = ''

    cidade_val = all_input_data.get('Cidade', '').strip()
    estado = all_input_data.get('Estado', '').strip()
    if cidade_val and estado:
        all_input_data['CIDADE_ESTADO'] = f"{cidade_val} - {estado}"
    elif cidade_val:
        all_input_data['CIDADE_ESTADO'] = cidade_val
    elif estado:
        all_input_data['CIDADE_ESTADO'] = estado
    else:
        all_input_data['CIDADE_ESTADO'] = ''

    endereco_rua_numero = all_input_data.get('Endereco_Rua_Numero', '').strip()
    bairro_val = all_input_data.get('Bairro', '').strip()
    if endereco_rua_numero and bairro_val:
        all_input_data['Endereco_Rua_Numero_Bairro'] = f"{endereco_rua_numero} - {bairro_val}"
    elif endereco_rua_numero:
        all_input_data['Endereco_Rua_Numero_Bairro'] = endereco_rua_numero
    elif bairro_val:
        all_input_data['Endereco_Rua_Numero_Bairro'] = bairro_val
    else:
        all_input_data['Endereco_Rua_Numero_Bairro'] = ''

    try:
        lat_dec = float(str(all_input_data.get('LATITUDE', '0')).replace(',', '.'))
        lon_dec = float(str(all_input_data.get('LONGITUDE', '0')).replace(',', '.'))
        all_input_data['LATITUDE_GMS'] = decimal_to_dms(lat_dec, is_latitude=True)
        all_input_data['LONGITUDE_GMS'] = decimal_to_dms(lon_dec, is_latitude=False)
    except (ValueError, TypeError):
        all_input_data['LATITUDE_GMS'] = 'N/A'
        all_input_data['LONGITUDE_GMS'] = 'N/A'

    data_hoje = datetime.now()
    all_input_data['DATA_ATUAL'] = data_hoje.strftime('%d/%m/%Y')
    
    if all_input_data.get('DATA_ART') == 'Não informado' or not all_input_data.get('DATA_ART'):
        all_input_data['DATA_ART'] = all_input_data['DATA_ATUAL']

    data_operacao_prevista = data_hoje + timedelta(days=90)
    all_input_data['DATA_OPERACAO_PREVISTA'] = data_operacao_prevista.strftime('%d/%m/%Y')

    try:
        potencia_wp_str = str(all_input_data.get('POTENCIA_MODULO_MANUAL', '0')).replace('Wp', '').strip().replace(',', '.')
        potencia_wp = float(potencia_wp_str)
        all_input_data['POTENCIA_MODULO_MANUAL_KWP'] = potencia_wp / 1000
    except ValueError:
        all_input_data['POTENCIA_MODULO_MANUAL_KWP'] = 'Não informado'

    session['temp_files_to_zip'] = []
    temp_zip_dir = tempfile.mkdtemp()

    nome_razao_social_clean = re.sub(r'[\/:*?"<>|]', '', all_input_data.get('Nome_Razao_Social', 'Cliente')).strip()
    if not nome_razao_social_clean:
        nome_razao_social_clean = 'Cliente'
    session['nome_razao_social_zip_folder'] = nome_razao_social_clean


    current_excel_fv_template_path = EXCEL_PROJETO_FV_TEMPLATE_PATH
    current_certidao_registro_profissional_path = CERTIDAO_REGISTRO_PROFISSIONAL_PATH
    
    temp_excel_proj_fv_filename_unique = f"dados_do_projeto_{uuid.uuid4().hex}.xlsx"
    temp_excel_proj_fv_filepath = os.path.join(temp_zip_dir, temp_excel_proj_fv_filename_unique)

    try:
        workbook_proj_fv = openpyxl.load_workbook(current_excel_fv_template_path)
        if 'DADOS' not in workbook_proj_fv.sheetnames:
            return render_template('index.html', error="Erro: A aba 'DADOS' não foi encontrada no arquivo Excel 'Planilha Projetos FV.xlsx'.")

        sheet_dados_proj_fv = workbook_proj_fv['DADOS']

        for key, cell_address in EXCEL_PROJETO_FV_CELL_MAPPING.items():
            value_to_write = all_input_data.get(key, '')

            if key == 'Endereco_Rua':
                full_address = all_input_data.get('Endereco_Rua_Numero', '')
                street, _ = parse_address_for_excel(str(full_address))
                sheet_dados_proj_fv[cell_address] = street
                continue
            elif key == 'Numero_Endereco':
                full_address = all_input_data.get('Endereco_Rua_Numero', '')
                _, number = parse_address_for_excel(str(full_address))
                sheet_dados_proj_fv[cell_address] = number
                continue

            if key in NUMERIC_KEYS_FOR_FORMATTING and value_to_write != '':
                try:
                    num_value = float(str(value_to_write).replace(',', '.'))
                    sheet_dados_proj_fv[cell_address] = num_value
                except ValueError:
                    sheet_dados_proj_fv[cell_address] = str(value_to_write)
            elif key in DATE_KEYS_FOR_FORMATTING and value_to_write != '':
                try:
                    dt_obj = datetime.strptime(str(value_to_write), '%Y-%m-%d')
                except ValueError:
                    try:
                        dt_obj = datetime.strptime(str(value_to_write), '%d/%m/%Y')
                    except ValueError:
                        try:
                            dt_obj = datetime.strptime(str(value_to_write), '%d-%m-%Y')
                        except ValueError:
                            dt_obj = None
                if dt_obj:
                    sheet_dados_proj_fv[cell_address] = dt_obj
                    sheet_dados_proj_fv[cell_address].number_format = 'DD/MM/YYYY'
                else:
                    sheet_dados_proj_fv[cell_address] = str(value_to_write)
            else:
                sheet_dados_proj_fv[cell_address] = str(value_to_write)

        workbook_proj_fv.save(temp_excel_proj_fv_filepath)
        session['temp_files_to_zip'].append({'path': temp_excel_proj_fv_filepath, 'zip_filename': 'Dados do Projeto.xlsx'})

    except FileNotFoundError:
        return render_template('index.html', error=f"Erro: Arquivo Excel de template '{EXCEL_PROJETO_FV_TEMPLATE_FILENAME}' não encontrado.")
    except Exception as e:
        return render_template('index.html', error=f"Erro ao processar/salvar dados no Excel de Projeto FV: {e}")

    temp_art_filename_unique = f"art_preenchida_{uuid.uuid4().hex}.txt"
    temp_art_filepath = os.path.join(temp_zip_dir, temp_art_filename_unique)
    try:
        art_content = generate_art_txt_content(all_input_data)
        with open(temp_art_filepath, 'w', encoding='utf-8') as f:
            f.write(art_content)
        session['temp_files_to_zip'].append({'path': temp_art_filepath, 'zip_filename': '2. ART - PREENCHIDO.txt'})
    except Exception as e:
        return render_template('index.html', error=f"Erro ao gerar ART.txt: {e}")
    
    image_files = {
        'foto_disjuntor': request.files.get('foto_disjuntor'),
        'foto_fachada': request.files.get('foto_fachada'),
        'foto_entrada_energia': request.files.get('foto_entrada_energia')
    }

    image_paths_for_pdf = []
    titles_for_pdf = {
        'foto_disjuntor': 'Foto do Disjuntor',
        'foto_fachada': 'Foto da Fachada',
        'foto_entrada_energia': 'Foto da Entrada de Energia'
    }

    for key, file_obj in image_files.items():
        if file_obj and file_obj.filename != '' and allowed_file(file_obj.filename, ALLOWED_IMAGE_EXTENSIONS):
            image_filename = secure_filename(file_obj.filename)
            temp_image_filepath = os.path.join(temp_zip_dir, image_filename)
            file_obj.save(temp_image_filepath)
            image_paths_for_pdf.append({'path': temp_image_filepath, 'title': titles_for_pdf[key]})

    if image_paths_for_pdf:
        temp_image_pdf_filename_unique = f"fotos_geral_fachada_{uuid.uuid4().hex}.pdf"
        temp_image_pdf_filepath = os.path.join(temp_zip_dir, temp_image_pdf_filename_unique)
        generate_images_pdf(image_paths_for_pdf, temp_image_pdf_filepath)
        session['temp_files_to_zip'].append({'path': temp_image_pdf_filepath, 'zip_filename': '6. 7. Foto Geral da Entrada de Energia - Fachada.pdf'})
    else:
        temp_empty_image_pdf_placeholder = os.path.join(temp_zip_dir, f"placeholder_fotos_vazio_geral_{uuid.uuid4().hex}.txt")
        with open(temp_empty_image_pdf_placeholder, 'w', encoding='utf-8') as f:
            f.write("Aviso: Nenhuma imagem da entrada elétrica (disjuntor, fachada, etc.) foi fornecida. Este arquivo é um placeholder.")
        session['temp_files_to_zip'].append({'path': temp_empty_image_pdf_placeholder, 'zip_filename': '6. 7. Foto Geral da Entrada de Energia - Fachada.txt'})

    session['temp_files_to_zip'].append({'path': current_certidao_registro_profissional_path, 'zip_filename': '1. Certidao de Registro Profissional.pdf'})

    empty_txt_files_common = {
        '9. 10. Documento de Identidade.txt': 'Conteúdo de placeholder para o documento de identidade.',
        '5. Certificado Inmetro Inversor Solar.txt': 'Conteúdo de placeholder para o certificado Inmetro do inversor solar.',
    }
    for zip_name, content in empty_txt_files_common.items():
        temp_empty_filepath = os.path.join(temp_zip_dir, f"empty_placeholder_{uuid.uuid4().hex}.txt")
        with open(temp_empty_filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        session['temp_files_to_zip'].append({'path': temp_empty_filepath, 'zip_filename': zip_name})

    if distributor_type == 'RGE':
        temp_postagem_filename_unique = f"postagem_{distributor_type.lower()}_{uuid.uuid4().hex}.txt"
        temp_postagem_filepath = os.path.join(temp_zip_dir, temp_postagem_filename_unique)
        try:
            postagem_content = generate_postagem_txt_content(all_input_data)
            with open(temp_postagem_filepath, 'w', encoding='utf-8') as f:
                f.write(postagem_content)
            session['temp_files_to_zip'].append({'path': temp_postagem_filepath, 'zip_filename': f'Postagem do projeto no site da {distributor_type}.txt'})
        except Exception as e:
            return render_template('index.html', error=f"Erro ao gerar Postagem_{distributor_type}.txt: {e}")

        current_anexo_f_template_path = ANEXO_F_TEMPLATE_PATH
        current_anexo_e_template_path = ANEXO_E_TEMPLATE_PATH
        current_termo_aceite_template_path = TERMO_ACEITE_TEMPLATE_PATH

        temp_excel_anexo_f_filename_unique = f"anexo_f_preenchido_{uuid.uuid4().hex}.xlsx"
        temp_excel_anexo_f_filepath = os.path.join(temp_zip_dir, temp_excel_anexo_f_filename_unique)

        try:
            workbook_anexo_f = openpyxl.load_workbook(current_anexo_f_template_path)
            if ANEXO_F_SHEET_NAME not in workbook_anexo_f.sheetnames:
                return render_template('index.html', error=f"Erro: A aba '{ANEXO_F_SHEET_NAME}' não foi encontrada no arquivo Excel '{ANEXO_F_TEMPLATE_FILENAME}'.")

            sheet_anexo_f = workbook_anexo_f[ANEXO_F_SHEET_NAME]

            for row in sheet_anexo_f.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        cell_text_standardized = cell.value.strip().upper()
                        if cell_text_standardized in ANEXO_F_PLACEHOLDER_TO_PYTHON_VAR:
                            python_var_name = ANEXO_F_PLACEHOLDER_TO_PYTHON_VAR[cell_text_standardized]
                            value_to_write = all_input_data.get(python_var_name)

                            if python_var_name == 'FULL_ADDRESS_COMPOSED':
                                rua_num_completo = all_input_data.get('Endereco_Rua_Numero', '')
                                bairro_comp = all_input_data.get('Bairro', '')
                                rua_parsed, numero_parsed = parse_address_for_excel(str(rua_num_completo))
                                full_address_str = f"{rua_parsed}"
                                if numero_parsed:
                                    full_address_str += f", {numero_parsed}"
                                if bairro_comp:
                                    full_address_str += f" - {bairro_comp}"
                                cell.value = full_address_str.strip(' ,-')
                            elif python_var_name in DATE_KEYS_FOR_FORMATTING:
                                if value_to_write:
                                    try:
                                        dt_obj = datetime.strptime(str(value_to_write), '%Y-%m-%d')
                                    except ValueError:
                                        try:
                                            dt_obj = datetime.strptime(str(value_to_write), '%d/%m/%Y')
                                        except ValueError:
                                            dt_obj = None
                                    if dt_obj:
                                        cell.number_format = 'DD/MM/YYYY'
                                        cell.value = dt_obj
                                    else:
                                        cell.value = str(value_to_write)
                                else:
                                    cell.value = ''
                            elif python_var_name in NUMERIC_KEYS_FOR_FORMATTING:
                                # **NOVO TRATAMENTO: LATITUDE e LONGITUDE como string formatada com ponto e 6 casas decimais**
                                if python_var_name in ['LATITUDE', 'LONGITUDE']:
                                    if value_to_write is not None:
                                        try:
                                            # Converte para float primeiro, para garantir o valor numérico
                                            float_val = float(str(value_to_write).replace(',', '.'))
                                            # Formata como string com 6 casas decimais e ponto
                                            cell.value = f"{float_val:.6f}"
                                        except ValueError:
                                            # Se a conversão falhar, mantém o valor original como string
                                            cell.value = str(value_to_write)
                                    else:
                                        cell.value = '' # Se for None
                                else:
                                    # Para os outros campos numéricos, mantém a formatação existente (2 casas, vírgula)
                                    cell.value = format_value_for_display(value_to_write, is_numeric=True)
                            else:
                                cell.value = str(value_to_write) if value_to_write is not None else ''

                            original_font = cell.font
                            cell.font = openpyxl.styles.Font(color='00000000', name=original_font.name, size=original_font.size)

            workbook_anexo_f.save(temp_excel_anexo_f_filepath)
            session['temp_files_to_zip'].append({'path': temp_excel_anexo_f_filepath, 'zip_filename': '3. Formulário Anexo F - PREENCHIDO.xlsx'})

        except FileNotFoundError:
            return render_template('index.html', error=f"Erro: Arquivo Excel de template '{ANEXO_F_TEMPLATE_FILENAME}' não encontrado.")
        except Exception as e:
            return render_template('index.html', error=f"Erro ao processar/salvar dados no Anexo F: {e}")

        temp_anexo_e_filename_unique = f"anexo_e_preenchido_{uuid.uuid4().hex}.docx"
        temp_anexo_e_filepath = os.path.join(temp_zip_dir, temp_anexo_e_filename_unique)

        try:
            docx_replacements_anexo_e = {
                'UC': get_formatted_value_for_doc('UC', all_input_data),
                'ENDERECO_RUA_NUMERO': get_formatted_value_for_doc('Endereco_Rua_Numero', all_input_data),
                'Bairro': get_formatted_value_for_doc('Bairro', all_input_data),
                'Cidade': get_formatted_value_for_doc('Cidade', all_input_data),
                'ESTADO': get_formatted_value_for_doc('Estado', all_input_data),
                'CEP': get_formatted_value_for_doc('CEP', all_input_data),
                'TELEFONE': get_formatted_value_for_doc('TELEFONE', all_input_data),
                'E-MAIL': get_formatted_value_for_doc('E-MAIL', all_input_data),
                'CARGA_INSTALADA': get_formatted_value_for_doc('CARGA_INSTALADA', all_input_data),
                'CATEGORIA': get_formatted_value_for_doc('CATEGORIA', all_input_data),
                'CLASSE_TARIFARIA': get_formatted_value_for_doc('Classe_Tarifaria', all_input_data),
                'GRUPO_TARIFARIO': get_formatted_value_for_doc('Grupo_Tarifario', all_input_data),
                'TENSAO_NOMINAL_V': get_formatted_value_for_doc('Tensao_Nominal_V', all_input_data),
                'NUMERO_FASES_CALCULADO': get_formatted_value_for_doc('NUMERO_FASES_CALCULADO', all_input_data),
                'POTENCIA_TOTAL_SISTEMA_KWP': get_formatted_value_for_doc('POTENCIA_TOTAL_SISTEMA_KWP', all_input_data),
                'FABRICANTE_INVERSOR_MANUAL': get_formatted_value_for_doc('FABRICANTE_INVERSOR_MANUAL', all_input_data),
                'MODELO_INVERSOR_CALCULADO': get_formatted_value_for_doc('MODELO_INVERSOR_CALCULADO', all_input_data),
                'POTENCIA_INVERSOR_MANUAL': get_formatted_value_for_doc('POTENCIA_INVERSOR_MANUAL', all_input_data),
                'DISJUNTOR_CA': get_formatted_value_for_doc('DISJUNTOR_CA', all_input_data),
                'DISJ_CA_INTR': get_formatted_value_for_doc('DISJ_CA_INTR', all_input_data),
                'DISJ_CA_TENS': get_formatted_value_for_doc('DISJ_CA_TENS', all_input_data),
                'DISJ_CA_ATEN': get_formatted_value_for_doc('DISJ_CA_ATEN', all_input_data),
                'ISOLACAO_CA': get_formatted_value_for_doc('ISOLACAO_CA', all_input_data),
                'CABO_CA': get_formatted_value_for_doc('CABO_CA', all_input_data),
                'DATA_ATUAL': get_formatted_value_for_doc('DATA_ATUAL', all_input_data),
                'POTENCIA_PICO_MODULOS': get_formatted_value_for_doc('POTENCIA_PICO_MODULOS', all_input_data),
                'POTENCIA_MODULO_MANUAL_KWP': get_formatted_value_for_doc('POTENCIA_MODULO_MANUAL_KWP', all_input_data),
                'QUANTIDADE_INVERSOR_MANUAL': get_formatted_value_for_doc('QUANTIDADE_INVERSOR_MANUAL', all_input_data),
                'Nome_Razao_Social': get_formatted_value_for_doc('Nome_Razao_Social', all_input_data), 
            }

            doc_anexo_e_preenchido = replace_docx_placeholders(current_anexo_e_template_path, docx_replacements_anexo_e)
            doc_anexo_e_preenchido.save(temp_anexo_e_filepath)
            session['temp_files_to_zip'].append({'path': temp_anexo_e_filepath, 'zip_filename': '11. Formulário Anexo E - PREENCHIDO.docx'})

        except FileNotFoundError:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_anexo_e_template_nao_encontrado_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Anexo E: Template '{ANEXO_E_TEMPLATE_FILENAME}' não encontrado.")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '11. Formulário Anexo E - ERRO (Template não encontrado).txt'})
        except Exception as e:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_anexo_e_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Anexo E: {e}")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '11. Formulário Anexo E - ERRO.txt'})

        temp_termo_aceite_filename_unique = f"termo_aceite_preenchido_{uuid.uuid4().hex}.docx"
        temp_termo_aceite_filepath = os.path.join(temp_zip_dir, temp_termo_aceite_filename_unique)

        try:
            docx_replacements_termo_aceite = {
                'Nome_Razao_Social': get_formatted_value_for_doc('Nome_Razao_Social', all_input_data), 
                'CPF_CNPJ': get_formatted_value_for_doc('CNPJ_CPF', all_input_data),
                'UC': get_formatted_value_for_doc('UC', all_input_data),
                'DATA_ATUAL': get_formatted_value_for_doc('DATA_ATUAL', all_input_data),
                'NUMERO_ART': get_formatted_value_for_doc('ART', all_input_data),
                'CIDADE': get_formatted_value_for_doc('Cidade', all_input_data),
                'ESTADO': get_formatted_value_for_doc('Estado', all_input_data),
            }

            doc_termo_aceite_preenchido = replace_docx_placeholders(current_termo_aceite_template_path, docx_replacements_termo_aceite)
            doc_termo_aceite_preenchido.save(temp_termo_aceite_filepath)
            session['temp_files_to_zip'].append({'path': temp_termo_aceite_filepath, 'zip_filename': '8. Termo de Aceite - PREENCHIDO.docx'})

        except FileNotFoundError:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_termo_aceite_template_nao_encontrado_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Termo de Aceite: Template '{TERMO_ACEITE_TEMPLATE_FILENAME}' não encontrado.")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '8. Termo de Aceite - ERRO (Template não encontrado).txt'})
        except Exception as e:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_termo_aceite_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao processar/salvar dados no Termo de Aceite: {e}")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '8. Termo de Aceite - ERRO.txt'})

        temp_empty_filepath = os.path.join(temp_zip_dir, f"empty_placeholder_projeto_{uuid.uuid4().hex}.txt")
        with open(temp_empty_filepath, 'w', encoding='utf-8') as f:
            f.write('Conteúdo de placeholder para o projeto.')
        session['temp_files_to_zip'].append({'path': temp_empty_filepath, 'zip_filename': '4. Projeto.txt'})

    elif distributor_type in ['COOPERLUZ', 'CERTHIL', 'CERMISSOES']:
        current_anexo_i_template_path = ANEXO_I_TEMPLATE_PATH
        current_procuracao_template_path = PROCURACAO_TEMPLATE_PATH
        current_termo_aceite_inciso_iii_template_path = TERMO_ACEITE_INCISO_III_TEMPLATE_PATH
        current_responsabilidade_tecnica_template_path = RESPONSABILIDADE_TECNICA_TEMPLATE_PATH
        current_dados_gd_ufv_template_path = DADOS_GD_UFV_TEMPLATE_PATH
        current_memorial_descritivo_template_path = MEMORIAL_DESCRITIVO_TEMPLATE_PATH

        temp_anexo_i_filename_unique = f"anexo_i_preenchido_{uuid.uuid4().hex}.xlsx"
        temp_anexo_i_filepath = os.path.join(temp_zip_dir, temp_anexo_i_filename_unique)

        try:
            workbook_anexo_i = openpyxl.load_workbook(current_anexo_i_template_path)
            if ANEXO_I_SHEET_NAME not in workbook_anexo_i.sheetnames:
                temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_anexo_i_template_nao_encontrado_{uuid.uuid4().hex}.txt")
                with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                    f.write(f"Erro ao gerar Anexo I: Template '{ANEXO_I_TEMPLATE_FILENAME}' não encontrado.")
                session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '4. Anexo I - ERRO (Template não encontrado).txt'})
            else:
                sheet_anexo_i = workbook_anexo_i[ANEXO_I_SHEET_NAME]

                for row in sheet_anexo_i.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            cell_text_standardized = cell.value.strip().upper()
                            if cell_text_standardized in ANEXO_I_PLACEHOLDER_TO_PYTHON_VAR:
                                python_var_name = ANEXO_I_PLACEHOLDER_TO_PYTHON_VAR[cell_text_standardized]
                                value_to_write = all_input_data.get(python_var_name)

                                if python_var_name in DATE_KEYS_FOR_FORMATTING:
                                    if value_to_write:
                                        try:
                                            dt_obj = datetime.strptime(value_to_write, '%d/%m/%Y')
                                            cell.number_format = 'DD/MM/YYYY'
                                            cell.value = dt_obj
                                        except ValueError:
                                            cell.value = str(value_to_write)
                                    else:
                                        cell.value = ''
                                elif python_var_name in NUMERIC_KEYS_FOR_FORMATTING:
                                    cell.value = format_value_for_display(value_to_write, is_numeric=True)
                                else:
                                    cell.value = str(value_to_write) if value_to_write is not None else ''

                                original_font = cell.font
                                cell.font = openpyxl.styles.Font(color='00000000', name=original_font.name, size=original_font.size)
                
                workbook_anexo_i.save(temp_anexo_i_filepath)
                session['temp_files_to_zip'].append({'path': temp_anexo_i_filepath, 'zip_filename': '4. Anexo I - PREENCHIDO.xlsx'})

        except FileNotFoundError:
            return render_template('index.html', error=f"Erro: Arquivo Excel de template '{ANEXO_I_TEMPLATE_FILENAME}' não encontrado.")
        except Exception as e:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_anexo_i_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Anexo I: {e}")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '4. Anexo I - ERRO.txt'})
        
        full_address_outorgante_parts = []
        rua_num_out = all_input_data.get('Endereco_Rua_Numero', '').strip()
        bairro_out = all_input_data.get('Bairro', '').strip()
        cidade_est_out = all_input_data.get('CIDADE_ESTADO', '').strip()
        cep_out = all_input_data.get('CEP', '').strip()

        if rua_num_out:
            full_address_outorgante_parts.append(rua_num_out)
        if bairro_out:
            full_address_outorgante_parts.append(f"– {bairro_out}")
        if cidade_est_out:
            full_address_outorgante_parts.append(f"na cidade de {cidade_est_out}")
        if cep_out:
             full_address_outorgante_parts.append(f"CEP {cep_out}")

        all_input_data['ENDERECO_COMPLETO_OUTORGANTE'] = ", ".join(filter(None, full_address_outorgante_parts))

        cidade_estado_data_str = f"{format_value_for_display(all_input_data.get('Cidade', '')).upper()} – {format_value_for_display(all_input_data.get('Estado', '')).upper()}, {get_formatted_value_for_doc('DATA_ATUAL', all_input_data)}"
        all_input_data['CIDADE_ESTADO_DATA_ASSINATURA'] = cidade_estado_data_str

        temp_procuracao_filename_unique = f"procuracao_preenchida_{uuid.uuid4().hex}.docx"
        temp_procuracao_filepath = os.path.join(temp_zip_dir, temp_procuracao_filename_unique)

        try:
            docx_replacements_procuracao = {
                'NOME_RAZAO_SOCIAL': get_formatted_value_for_doc('Nome_Razao_Social', all_input_data),
                'CPF_CNPJ': get_formatted_value_for_doc('CNPJ_CPF', all_input_data),
                'ENDERECO_RUA_NUMERO': get_formatted_value_for_doc('Endereco_Rua_Numero', all_input_data),
                'BAIRRO': get_formatted_value_for_doc('Bairro', all_input_data),
                'CIDADE': get_formatted_value_for_doc('Cidade', all_input_data),
                'ESTADO': get_formatted_value_for_doc('Estado', all_input_data),
                'UC': get_formatted_value_for_doc('UC', all_input_data),
                'DATA_ATUAL': get_formatted_value_for_doc('DATA_ATUAL', all_input_data),
                'ENDERECO_COMPLETO_OUTORGANTE': get_formatted_value_for_doc('ENDERECO_COMPLETO_OUTORGANTE', all_input_data),
                'CIDADE_ESTADO_DATA_ASSINATURA': get_formatted_value_for_doc('CIDADE_ESTADO_DATA_ASSINATURA', all_input_data),
            }

            doc_procuracao_preenchida = replace_docx_placeholders(current_procuracao_template_path, docx_replacements_procuracao)
            doc_procuracao_preenchida.save(temp_procuracao_filepath)
            session['temp_files_to_zip'].append({'path': temp_procuracao_filepath, 'zip_filename': '12. Procuração - PREENCHIDO.docx'})

        except FileNotFoundError:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_procuracao_template_nao_encontrado_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Procuração: Template '{PROCURACAO_TEMPLATE_FILENAME}' não encontrado.")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '12. Procuração - ERRO (Template não encontrado).txt'})
        except Exception as e:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_procuracao_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Procuração: {e}")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '12. Procuração - ERRO.txt'})

        temp_termo_aceite_inciso_iii_filename_unique = f"termo_aceite_inciso_iii_preenchido_{uuid.uuid4().hex}.docx"
        temp_termo_aceite_inciso_iii_filepath = os.path.join(temp_zip_dir, temp_termo_aceite_inciso_iii_filename_unique)

        try:
            docx_replacements_termo_aceite_inciso_iii = {
                'NOME_RAZAO_SOCIAL': get_formatted_value_for_doc('Nome_Razao_Social', all_input_data),
                'CPF_CNPJ': get_formatted_value_for_doc('CNPJ_CPF', all_input_data),
                'UC': get_formatted_value_for_doc('UC', all_input_data),
                'DATA_ATUAL': get_formatted_value_for_doc('DATA_ATUAL', all_input_data),
            }

            doc_termo_aceite_inciso_iii_preenchido = replace_docx_placeholders(current_termo_aceite_inciso_iii_template_path, docx_replacements_termo_aceite_inciso_iii)
            doc_termo_aceite_inciso_iii_preenchido.save(temp_termo_aceite_inciso_iii_filepath)
            session['temp_files_to_zip'].append({'path': temp_termo_aceite_inciso_iii_filepath, 'zip_filename': '13. Termo de Aceite Inciso III - PREENCHIDO.docx'})

        except FileNotFoundError:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_termo_aceite_inciso_iii_template_nao_encontrado_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Termo de Aceite Inciso III: Template '{TERMO_ACEITE_INCISO_III_TEMPLATE_FILENAME}' não encontrado.")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '13. Termo de Aceite Inciso III - ERRO (Template não encontrado).txt'})
        except Exception as e:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_termo_aceite_inciso_iii_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Termo de Aceite Inciso III: {e}")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '13. Termo de Aceite Inciso III - ERRO.txt'})
        
        temp_responsabilidade_tecnica_filename_unique = f"responsabilidade_tecnica_preenchido_{uuid.uuid4().hex}.docx"
        temp_responsabilidade_tecnica_filepath = os.path.join(temp_zip_dir, temp_responsabilidade_tecnica_filename_unique)

        try:
            docx_replacements_responsabilidade_tecnica = {
                'POTENCIA_TOTAL_SISTEMA_KWP': get_formatted_value_for_doc('POTENCIA_TOTAL_SISTEMA_KWP', all_input_data),
                'Tensao_Nominal_V': get_formatted_value_for_doc('Tensao_Nominal_V', all_input_data),
                'NUMERO_FASES_CALCULADO': get_formatted_value_for_doc('NUMERO_FASES_CALCULADO', all_input_data),
                'UC': get_formatted_value_for_doc('UC', all_input_data),
                'NOME_RAZAO_SOCIAL': get_formatted_value_for_doc('Nome_Razao_Social', all_input_data),
                'ENDERECO_RUA_NUMERO': get_formatted_value_for_doc('Endereco_Rua_Numero', all_input_data),
                'BAIRRO': get_formatted_value_for_doc('Bairro', all_input_data),
                'CIDADE': get_formatted_value_for_doc('Cidade', all_input_data),
                'ESTADO': get_formatted_value_for_doc('Estado', all_input_data),
                'ART': get_formatted_value_for_doc('ART', all_input_data),
                'MODELO_INVERSOR_CALCULADO': get_formatted_value_for_doc('MODELO_INVERSOR_CALCULADO', all_input_data),
                'POTENCIA_INVERSOR_MANUAL': get_formatted_value_for_doc('POTENCIA_INVERSOR_MANUAL', all_input_data),
                'DATA_ATUAL': get_formatted_value_for_doc('DATA_ATUAL', all_input_data),
                'INMETRO': get_formatted_value_for_doc('INMETRO', all_input_data),
            }

            doc_responsabilidade_tecnica_preenchido = replace_docx_placeholders(current_responsabilidade_tecnica_template_path, docx_replacements_responsabilidade_tecnica)
            doc_responsabilidade_tecnica_preenchido.save(temp_responsabilidade_tecnica_filepath)
            session['temp_files_to_zip'].append({'path': temp_responsabilidade_tecnica_filepath, 'zip_filename': '14. Responsabilidade Tecnica - PREENCHIDO.docx'})

        except FileNotFoundError:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_responsabilidade_tecnica_template_nao_encontrado_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Responsabilidade Tecnica: Template '{RESPONSABILIDADE_TECNICA_TEMPLATE_FILENAME}' não encontrado.")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '14. Responsabilidade Tecnica - ERRO (Template não encontrado).txt'})
        except Exception as e:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_responsabilidade_tecnica_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Responsabilidade Tecnica: {e}")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '14. Responsabilidade Tecnica - ERRO.txt'})

        temp_dados_gd_ufv_filename_unique = f"dados_gd_ufv_preenchido_{uuid.uuid4().hex}.docx"
        temp_dados_gd_ufv_filepath = os.path.join(temp_zip_dir, temp_dados_gd_ufv_filename_unique)

        try:
            docx_replacements_dados_gd_ufv = {
                'Classe_Tarifaria': get_formatted_value_for_doc('Classe_Tarifaria', all_input_data),
                'Grupo_Tarifario': get_formatted_value_for_doc('Grupo_Tarifario', all_input_data),
                'Cidade': get_formatted_value_for_doc('Cidade', all_input_data),
                'Estado': get_formatted_value_for_doc('Estado', all_input_data),
                'Endereco_Rua_Numero': get_formatted_value_for_doc('Endereco_Rua_Numero', all_input_data),
                'Bairro': get_formatted_value_for_doc('Bairro', all_input_data),
                'CEP': get_formatted_value_for_doc('CEP', all_input_data),
                'LATITUDE_GMS': get_formatted_value_for_doc('LATITUDE_GMS', all_input_data),
                'LONGITUDE_GMS': get_formatted_value_for_doc('LONGITUDE_GMS', all_input_data),
                'CNPJ_CPF': get_formatted_value_for_doc('CNPJ_CPF', all_input_data),
                'Nome_Razao_Social': get_formatted_value_for_doc('Nome_Razao_Social', all_input_data),
                'TELEFONE': get_formatted_value_for_doc('TELEFONE', all_input_data),
                'E-MAIL': get_formatted_value_for_doc('E-MAIL', all_input_data),
                'POTENCIA_TOTAL_SISTEMA_KWP': get_formatted_value_for_doc('POTENCIA_TOTAL_SISTEMA_KWP', all_input_data),
                'QUANTIDADE_PLACAS_MANUAL': get_formatted_value_for_doc('QUANTIDADE_PLACAS_MANUAL', all_input_data),
                'FABRICANTE_MODULO_MANUAL': get_formatted_value_for_doc('FABRICANTE_MODULO_MANUAL', all_input_data),
                'MODELO_MODULO_CALCULADO': get_formatted_value_for_doc('MODELO_MODULO_CALCULADO', all_input_data),
                'POTENCIA_INVERSOR_MANUAL': get_formatted_value_for_doc('POTENCIA_INVERSOR_MANUAL', all_input_data),
                'QUANTIDADE_INVERSOR_MANUAL': get_formatted_value_for_doc('QUANTIDADE_INVERSOR_MANUAL', all_input_data),
                'FABRICANTE_INVERSOR_MANUAL': get_formatted_value_for_doc('FABRICANTE_INVERSOR_MANUAL', all_input_data),
                'MODELO_INVERSOR_CALCULADO': get_formatted_value_for_doc('MODELO_INVERSOR_CALCULADO', all_input_data),
                'AREA': get_formatted_value_for_doc('AREA_ARRANJOS_CALCULADO', all_input_data),
                'DATA_OPERACAO_PREVISTA': get_formatted_value_for_doc('DATA_OPERACAO_PREVISTA', all_input_data),
                'POTENCIA_PICO_MODULOS': get_formatted_value_for_doc('POTENCIA_PICO_MODULOS', all_input_data),
                'POTENCIA_MODULO_MANUAL_KWP': get_formatted_value_for_doc('POTENCIA_MODULO_MANUAL_KWP', all_input_data),
            }

            doc_dados_gd_ufv_preenchido = replace_docx_placeholders(current_dados_gd_ufv_template_path, docx_replacements_dados_gd_ufv)
            doc_dados_gd_ufv_preenchido.save(temp_dados_gd_ufv_filepath)
            session['temp_files_to_zip'].append({'path': temp_dados_gd_ufv_filepath, 'zip_filename': '15. Dados para GD de UFV - PREENCHIDO.docx'})

        except FileNotFoundError:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_dados_gd_ufv_template_nao_encontrado_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Dados para GD de UFV: Template '{DADOS_GD_UFV_TEMPLATE_FILENAME}' não encontrado.")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '15. Dados para GD de UFV - ERRO (Template não encontrado).txt'})
        except Exception as e:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_dados_gd_ufv_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Dados para GD de UFV: {e}")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '15. Dados para GD de UFV - ERRO.txt'})

        temp_memorial_descritivo_filename_unique = f"memorial_descritivo_preenchido_{uuid.uuid4().hex}.docx"
        temp_memorial_descritivo_filepath = os.path.join(temp_zip_dir, temp_memorial_descritivo_filename_unique)

        try:
            docx_replacements_memorial_descritivo = {
                'Nome_Razao_Social': get_formatted_value_for_doc('Nome_Razao_Social', all_input_data),
                'UC': get_formatted_value_for_doc('UC', all_input_data),
                'Endereco_Rua_Numero': get_formatted_value_for_doc('Endereco_Rua_Numero', all_input_data),
                'Bairro': get_formatted_value_for_doc('Bairro', all_input_data),
                'Cidade': get_formatted_value_for_doc('Cidade', all_input_data),
                'Estado': get_formatted_value_for_doc('Estado', all_input_data),
                'LATITUDE_GMS': get_formatted_value_for_doc('LATITUDE_GMS', all_input_data),
                'LONGITUDE_GMS': get_formatted_value_for_doc('LONGITUDE_GMS', all_input_data),
                'ART': get_formatted_value_for_doc('ART', all_input_data),
                'QUANTIDADE_PLACAS_MANUAL': get_formatted_value_for_doc('QUANTIDADE_PLACAS_MANUAL', all_input_data),
                'FABRICANTE_MODULO_MANUAL': get_formatted_value_for_doc('FABRICANTE_MODULO_MANUAL', all_input_data),
                'MODELO_MODULO_CALCULADO': get_formatted_value_for_doc('MODELO_MODULO_CALCULADO', all_input_data),
                'POTENCIA_TOTAL_SISTEMA_KWP': get_formatted_value_for_doc('POTENCIA_TOTAL_SISTEMA_KWP', all_input_data),
                'QUANTIDADE_INVERSOR_MANUAL': get_formatted_value_for_doc('QUANTIDADE_INVERSOR_MANUAL', all_input_data),
                'FABRICANTE_INVERSOR_MANUAL': get_formatted_value_for_doc('FABRICANTE_INVERSOR_MANUAL', all_input_data),
                'MODELO_INVERSOR_CALCULADO': get_formatted_value_for_doc('MODELO_INVERSOR_CALCULADO', all_input_data),
                'POTENCIA_INVERSOR_MANUAL': get_formatted_value_for_doc('POTENCIA_INVERSOR_MANUAL', all_input_data),
                'DISJUNTOR_CA': get_formatted_value_for_doc('DISJUNTOR_CA', all_input_data),
                'DISJ_CA_INTR': get_formatted_value_for_doc('DISJ_CA_INTR', all_input_data),
                'DISJ_CA_TENS': get_formatted_value_for_doc('DISJ_CA_TENS', all_input_data),
                'DISJ_CA_ATEN': get_formatted_value_for_doc('DISJ_CA_ATEN', all_input_data),
                'ISOLACAO_CA': get_formatted_value_for_doc('ISOLACAO_CA', all_input_data),
                'CABO_CA': get_formatted_value_for_doc('CABO_CA', all_input_data),
                'DATA_ATUAL': get_formatted_value_for_doc('DATA_ATUAL', all_input_data),
                'POTENCIA_PICO_MODULOS': get_formatted_value_for_doc('POTENCIA_PICO_MODULOS', all_input_data),
                'POTENCIA_MODULO_MANUAL_KWP': get_formatted_value_for_doc('POTENCIA_MODULO_MANUAL_KWP', all_input_data),
            }

            doc_memorial_descritivo_preenchido = replace_docx_placeholders(current_memorial_descritivo_template_path, docx_replacements_memorial_descritivo)
            doc_memorial_descritivo_preenchido.save(temp_memorial_descritivo_filepath)
            session['temp_files_to_zip'].append({'path': temp_memorial_descritivo_filepath, 'zip_filename': '16. Memorial Descritivo Padrão Fecoergs - PREENCHIDO.docx'})

        except FileNotFoundError:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_memorial_descritivo_template_nao_encontrado_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Memorial Descritivo: Template '{MEMORIAL_DESCRITIVO_TEMPLATE_FILENAME}' não encontrado.")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '16. Memorial Descritivo - ERRO (Template não encontrado).txt'})
        except Exception as e:
            temp_error_txt_filepath = os.path.join(temp_zip_dir, f"erro_memorial_descritivo_{uuid.uuid4().hex}.txt")
            with open(temp_error_txt_filepath, 'w', encoding='utf-8') as f:
                f.write(f"Erro ao gerar Memorial Descritivo: {e}")
            session['temp_files_to_zip'].append({'path': temp_error_txt_filepath, 'zip_filename': '16. Memorial Descritivo - ERRO.txt'})

        empty_txt_files_cooperluz_specific = {
            'Diagrama Unifilar indicando desde o ponto de conexão com a Cooperluz.txt': 'Conteúdo de placeholder para Diagrama Unifilar.',
            'Planta de Localizacao e Situacao.txt': 'Conteúdo de placeholder para Planta de Localização e Situação.',
            'Diagrama Multifilar indicando desde o ponto de conexão com a Cooperluz.txt': 'Conteúdo de placeholder para Diagrama Multifilar.',
        }
        for zip_name, content in empty_txt_files_cooperluz_specific.items():
            temp_empty_filepath = os.path.join(temp_zip_dir, f"empty_placeholder_{uuid.uuid4().hex}.txt")
            with open(temp_empty_filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            session['temp_files_to_zip'].append({'path': temp_empty_filepath, 'zip_filename': zip_name})

    session['temp_zip_dir'] = temp_zip_dir
    session['all_input_data_for_correction'] = all_input_data

    final_zip_filename_unique = f"Projeto de {nome_razao_social_clean}_{uuid.uuid4().hex}.zip"

    lat_val = all_input_data.get('LATITUDE')
    lon_val = all_input_data.get('LONGITUDE')
    
    display_latitude = 'Não informado'
    if lat_val is not None and str(lat_val).strip() != '':
        try:
            display_latitude = f"{float(str(lat_val).replace(',', '.')):.6f}"
        except ValueError:
            display_latitude = str(lat_val)
    
    display_longitude = 'Não informado'
    if lon_val is not None and str(lon_val).strip() != '':
        try:
            display_longitude = f"{float(str(lon_val).replace(',', '.')):.6f}"
        except ValueError:
            display_longitude = str(lon_val)

    dados_cliente_for_display = {
        'Nome/Razão Social': format_value_for_display(all_input_data.get('Nome_Razao_Social')),
        'CNPJ/CPF': format_value_for_display(all_input_data.get('CNPJ_CPF')),
        'E-Mail': format_value_for_display(all_input_data.get('E-MAIL')),
        'Telefone': format_value_for_display(all_input_data.get('TELEFONE')),
        'Endereço': format_value_for_display(all_input_data.get('Endereco_Rua_Numero')),
        'Bairro': format_value_for_display(all_input_data.get('Bairro')),
        'Cidade': format_value_for_display(all_input_data.get('Cidade')),
        'Estado': format_value_for_display(all_input_data.get('Estado')),
        'CEP': format_value_for_display(all_input_data.get('CEP')),
        'UC': format_value_for_display(all_input_data.get('UC')),
        'Classe Tarifária': format_value_for_display(all_input_data.get('Classe_Tarifaria')),
        'Grupo Tarifário': format_value_for_display(all_input_data.get('Grupo_Tarifario')),
        'Tensão Nominal (V)': format_value_for_display(all_input_data.get('Tensao_Nominal_V'), is_numeric=True), 
        'Latitude': display_latitude,
        'Longitude': display_longitude,
    }

    carga_instalada_display = format_value_for_display(all_input_data.get('CARGA_INSTALADA'), is_numeric=True)
    if carga_instalada_display != 'Não informado':
        carga_instalada_display += 'kW'

    dados_entrada_for_display = {
        'Categoria': format_value_for_display(all_input_data.get('CATEGORIA')),
        'Carga Instalada': carga_instalada_display,
        'Tipo de Atendimento': format_value_for_display(all_input_data.get('TIPO_DE_ATENDIMENTO')),
        'Tipo de Caixa': format_value_for_display(all_input_data.get('TIPO_DE_CAIXA')),
        'Isolação': format_value_for_display(all_input_data.get('ISOLACAO')),
        'Número de Fases': format_value_for_display(all_input_data.get('NUMERO_FASES_CALCULADO')),
        'Ramal de Entrada': format_value_for_display(all_input_data.get('RAMAL_ENTRADA_CALCULADO')),
        'Disjuntor': format_value_for_display(all_input_data.get('DISJUNTOR_CALCULADO'), is_numeric=True),
    }

    dados_sistema_fv_for_display = {
        'ART': format_value_for_display(all_input_data.get('ART')),
        'Data da ART': format_value_for_display(all_input_data.get('DATA_ART'), is_date=True, date_separator='-'),
        'INMETRO do Inversor': format_value_for_display(all_input_data.get('INMETRO')),
        'Quantidade de Placas': format_value_for_display(all_input_data.get('QUANTIDADE_PLACAS_MANUAL')),
        'Potência do Módulo (Wp)': format_value_for_display(all_input_data.get('POTENCIA_MODULO_MANUAL'), is_numeric=True),
        'Fabricante do Módulo': format_value_for_display(all_input_data.get('FABRICANTE_MODULO_MANUAL')),
        'Modelo do Módulo': format_value_for_display(all_input_data.get('MODELO_MODULO_CALCULADO')),
        'Área Total dos Arranjos (m²)': format_value_for_display(all_input_data.get('AREA_ARRANJOS_CALCULADO'), is_numeric=True),
        'Potência Pico dos Módulos (kW)': format_value_for_display(all_input_data.get('POTENCIA_PICO_MODULOS'), is_numeric=True),
        'Quantidade de Inversores': format_value_for_display(all_input_data.get('QUANTIDADE_INVERSOR_MANUAL')),
        'Potência do Inversor (kW)': format_value_for_display(all_input_data.get('POTENCIA_INVERSOR_MANUAL'), is_numeric=True),
        'Fabricante do Inversor': format_value_for_display(all_input_data.get('FABRICANTE_INVERSOR_MANUAL')),
        'Modelo do Inversor': format_value_for_display(all_input_data.get('MODELO_INVERSOR_CALCULADO')),
        'Potência Total do Sistema (kWp)': format_value_for_display(all_input_data.get('POTENCIA_TOTAL_SISTEMA_KWP'), is_numeric=True),
        'Isolação CA': format_value_for_display(all_input_data.get('ISOLACAO_CA')),
        'Cabo CA': format_value_for_display(all_input_data.get('CABO_CA')),
        'Disjuntor CA': format_value_for_display(all_input_data.get('DISJUNTOR_CA')),
    }

    return render_template('success.html',
                           message="Todos os documentos foram processados e estão prontos para download em um único arquivo ZIP!",
                           zip_filename=final_zip_filename_unique,
                           distributor_type=distributor_type,
                           dados_cliente=dados_cliente_for_display,
                           dados_entrada=dados_entrada_for_display,
                           dados_sistema_fv=dados_sistema_fv_for_display)


@app.route('/download_zip/<zip_filename>', methods=['GET'])
def download_zip(zip_filename):
    files_to_zip_info = session.pop('temp_files_to_zip', [])
    temp_zip_dir = session.pop('temp_zip_dir', None)
    nome_razao_social_zip_folder = session.pop('nome_razao_social_zip_folder', 'Projeto')

    if not files_to_zip_info or not temp_zip_dir:
        return "Erro: Arquivo ZIP não encontrado ou sessão expirada. Por favor, tente novamente.", 404

    zip_filepath = os.path.join(tempfile.gettempdir(), zip_filename)

    try:
        with zipfile.ZipFile(zip_filepath, 'w', zipfile.ZIP_DEFLATED) as zf:
            for file_info in files_to_zip_info:
                original_path = file_info['path']
                zip_name = file_info['zip_filename']
                zf.write(original_path, arcname=f"{nome_razao_social_zip_folder}/{zip_name}")

        response = send_file(zip_filepath,
                             mimetype='application/zip',
                             as_attachment=True,
                             download_name=f"Projeto de {nome_razao_social_zip_folder}.zip")

        @response.call_on_close
        def cleanup():
            try:
                if os.path.exists(zip_filepath):
                    os.remove(zip_filepath)
                if temp_zip_dir and os.path.exists(temp_zip_dir):
                    shutil.rmtree(temp_zip_dir)
            except Exception as e:
                print(f"Erro durante a limpeza de arquivos temporários: {e}")

        return response

    except Exception as e:
        return f"Erro ao gerar o arquivo ZIP: {e}", 500

if __name__ == '__main__':
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    app.run(host='0.0.0.0', port=5000, debug=True)
